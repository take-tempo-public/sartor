"""Document output generation — P1 Hardening.

Deterministic conversion of LLM-generated content into downloadable documents.
When the primary resume is .docx, it is opened as a style template so the
output inherits the original's fonts, margins, heading styles, and list
formatting exactly. For PDF or Markdown originals, clean defaults are used.
"""

import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.numbering import CT_NumPr
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Matches any common bullet prefix used by LLMs:
# -, *, •, –, —, ·, ◆, ●, ▪, ›, ‣
BULLET_RE = re.compile(r"^[-*\u2022\u2013\u2014\u00b7\u25c6\u25cf\u25aa\u2023\u2043\u203a]\s+")

# Inline markdown: ***bold+italic***, **bold**, *italic*
_INLINE_RE = re.compile(r"(\*\*\*[^*\n]+?\*\*\*|\*\*[^*\n]+?\*\*|\*[^*\n]+?\*)")

# Markdown structural markers used by _normalize_markdown to re-inject
# newlines when an LLM smushes a whole resume into a single line.
# Header boundary: only the FIRST `#` of an `# / ## / ###` heading (the
# non-`#` lookbehind prevents matching the 2nd `#` of `##` etc.).
_MD_HEADER_BOUNDARY_RE = re.compile(r"(?<![\n#])(?=#{1,3}\s)")
# Bullet boundary: `- <Capital>` preceded by text (letter / digit /
# period / question / exclamation / tab / closing paren). The capital-
# letter requirement avoids false positives on hyphenated words
# ("front-end", "real-time").
_MD_BULLET_BOUNDARY_RE = re.compile(r"(?<=[.!?\w\t)])(- [A-Z])")
# H2 → body boundary. After a `## <Title>` (single-word title, h2 only —
# `##(?!#)` excludes h3), break before the first body char (any capital).
# `\w+?` is non-greedy so it stops at the shortest word before the next
# capital letter — without this, `## SkillsUX` would split as
# `## SkillsU\nX` (greedy `\w+` consumes into the body word).
# Multi-word h2 titles ("Work Experience") are a known limitation; the
# LLM should emit them with proper newlines or the heading will be
# clipped at the first word.
_MD_H2_BODY_BOUNDARY_RE = re.compile(r"(##(?!#)\s+\w+?)([A-Z])")
_MD_TRIPLE_NEWLINE_RE = re.compile(r"\n{3,}")


def _normalize_markdown(content: str) -> str:
    """Re-inject newlines lost when an LLM emits markdown as one line.

    P1 Hardening — deterministic Python repairs LLM formatting drift so
    every downstream consumer (markdown export and `md_to_json_resume`,
    which the .docx / .pdf / HTML renderers all parse) sees the structural
    markers on their own lines. Conservative by design: only inserts
    newlines around unambiguous markdown markers, never inside running
    text.

    Markers handled:
      `# ` / `## ` / `### `       — start a new line preceded by a blank
      `- <Capital>`                — bullet preceded by text; +1 newline
      `## <Title><Body>`           — single-word section title → body

    Not handled (rare, left to LLM-side):
      The name / subtitle / contact triad smushed into the first line
      (e.g. `# Jane DoeSenior EngineerJane@ex.com`). There is no clean
      algorithmic signal for where each chunk ends — the LLM should
      emit them on separate lines. A prompt-side fix in
      `analyzer.py:RESUME_RULES` is the right lever; this normalizer
      handles the structural body of the document so a smushed top
      line is recoverable by hand.

    Idempotent: already-well-formed markdown passes through unchanged.
    """
    if not content:
        return content
    # Pass 1: blank line before headers
    content = _MD_HEADER_BOUNDARY_RE.sub("\n\n", content)
    # Pass 2: newline before bullets
    content = _MD_BULLET_BOUNDARY_RE.sub(r"\n\1", content)
    # Pass 3: blank line between h2 title and body
    content = _MD_H2_BODY_BOUNDARY_RE.sub(r"\1\n\n\2", content)
    # Pass 4: collapse 3+ newlines to 2
    content = _MD_TRIPLE_NEWLINE_RE.sub("\n\n", content)
    return content.strip() + "\n"


def generate_resume(
    content: str,
    output_format: str,
    username: str,
    base_dir: str = "output",
    template_path: str | None = None,
) -> str:
    """Generate the tailored resume in the requested format.

    Supported formats:
      .md   — write the normalized markdown directly
      .docx — render via python-docx + the persona .docx as style template
      .pdf  — render via Playwright (β.3): parse markdown → JSON Resume
              → Jinja2 HTML → headless Chromium → PDF. Requires the
              persona's `.html` + `.css` companions next to the .docx
              file, and the Chromium binary (one-time install via
              `python -m playwright install chromium`).

    Every call also writes a JSON Resume v1.0 sidecar
    (`resume_{ts}.jsonresume.json`) next to the primary output (β.2).
    The sidecar is best-effort — write failures log a warning but do
    not block the primary output.
    """
    import json as _json
    from pathlib import Path as _Path

    from json_resume import md_to_json_resume

    out_dir = Path(base_dir) / username
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Normalize before parsing — md_to_json_resume relies on the structural
    # markers (headings, bullets) sitting on their own lines; a smushed payload
    # would collapse into one giant section. Every renderer (.md, .docx, .pdf,
    # HTML preview) then flows from this single json_doc, so they cannot diverge.
    content = _normalize_markdown(content)
    json_doc = md_to_json_resume(content)

    if output_format == ".md":
        path = out_dir / f"resume_{ts}.md"
        path.write_text(content, encoding="utf-8")
    elif output_format == ".pdf":
        path = out_dir / f"resume_{ts}.pdf"
        _render_pdf_from_json(json_doc, template_path, path)
    else:
        path = out_dir / f"resume_{ts}.docx"
        # D3 (single source of truth): the .docx renders from the SAME
        # JSON Resume `json_doc` the preview/PDF use, so DOWNLOAD == PREVIEW.
        # (Previously the .docx parsed the raw markdown itself — a second,
        # divergent engine that emitted off-map `## headings` the preview
        # silently dropped and matched neither the preview nor the source.)
        _write_docx_from_json_resume(json_doc, path, template_path=template_path)

    # JSON Resume sidecar — best-effort. Parsing is deterministic and
    # tested; realistic failure modes are disk-write issues.
    try:
        sidecar = _Path(str(path)).with_suffix(".jsonresume.json")
        sidecar.write_text(_json.dumps(json_doc, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception as exc:
        import logging

        logging.getLogger(__name__).warning(
            "JSON Resume sidecar write failed for %s: %s",
            path,
            exc,
        )

    return str(path)


def generate_resume_from_json_resume(
    json_doc: dict[str, Any],
    output_format: str,
    username: str,
    base_dir: str = "output",
    template_path: str | None = None,
) -> str:
    """Render a pre-built JSON Resume doc to the requested format WITHOUT parsing markdown.

    Generation-experience re-architecture Phase 4: the corpus-mode deterministic
    assemble already holds the frozen ``approved_composition`` (a JSON Resume dict),
    so it renders that doc DIRECTLY through the same writers `generate_resume` uses
    after its `md_to_json_resume` parse (D3 single source of truth) — download ==
    preview == ``approved_composition`` by construction, with no LLM and no markdown
    round-trip. The ``.md`` format serializes the doc via `json_resume_to_markdown`.
    Writes the same ``.jsonresume.json`` sidecar. Returns the primary output path.
    """
    import json as _json
    from pathlib import Path as _Path

    out_dir = Path(base_dir) / username
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    if output_format == ".md":
        from json_resume import json_resume_to_markdown

        path = out_dir / f"resume_{ts}.md"
        path.write_text(json_resume_to_markdown(json_doc), encoding="utf-8")
    elif output_format == ".pdf":
        path = out_dir / f"resume_{ts}.pdf"
        _render_pdf_from_json(json_doc, template_path, path)
    else:
        path = out_dir / f"resume_{ts}.docx"
        _write_docx_from_json_resume(json_doc, path, template_path=template_path)

    try:
        sidecar = _Path(str(path)).with_suffix(".jsonresume.json")
        sidecar.write_text(_json.dumps(json_doc, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception as exc:
        import logging

        logging.getLogger(__name__).warning(
            "JSON Resume sidecar write failed for %s: %s", path, exc
        )

    return str(path)


def _render_pdf_from_json(
    json_doc: dict,
    docx_template_path: str | None,
    output_pdf_path: Path,
) -> None:
    """Render the JSON Resume document to PDF via the persona's HTML companion.

    The companion is the .html sibling of the .docx template
    (e.g. `personas/bundled/classic.docx` → `personas/bundled/classic.html`).

    Falls back to the bundled Classic HTML template if the resolved
    persona doesn't ship an .html companion yet — keeps the output
    path working as more personas pick up HTML companions over time.
    """
    from pdf_render import html_template_path_for, render_pdf

    html_template: Path | None = None
    if docx_template_path:
        html_template = html_template_path_for(docx_template_path)
        if html_template is None:
            # Lazily generate the companion so a PDF of an uploaded template
            # honors its own typography instead of falling back to Classic
            # (walkthrough B2 — mirrors the live-preview route). Deterministic.
            from docx_to_persona_html import generate_companion

            companion = generate_companion(docx_template_path)
            if companion is not None:
                html_template = companion[0]

    if html_template is None:
        # Fallback: bundled Classic. The path resolves relative to the
        # generator.py file at the repo root so it works whether the
        # caller passed an absolute or relative template_path.
        repo_root = Path(__file__).resolve().parent
        fallback = repo_root / "personas" / "bundled" / "classic.html"
        if fallback.exists():
            html_template = fallback

    if html_template is None:
        raise FileNotFoundError(
            "No HTML persona template available for PDF rendering. "
            "Drop classic.html + classic.css into personas/bundled/."
        )

    render_pdf(json_doc, html_template_path=html_template, output_pdf_path=output_pdf_path)


def generate_cover_letter(
    content: str,
    username: str,
    base_dir: str = "output",
    output_format: str = ".docx",
    template_path: str | None = None,
) -> str:
    """Generate the cover letter in the requested format.

    Supported formats (mirrors `generate_resume`):
      .md   — write the normalized markdown directly
      .docx — business-letter `.docx` via `_write_cover_letter_docx` (persona
              font from the chosen template, dense single-spaced body, no name
              banner, inline addressee — the 2026-05-26 styling decisions)
      .pdf  — render through the shared `personas/cover_letter.html` shell via
              Playwright (`_render_cover_letter_pdf`), byte-faithful to the
              Step-6 live preview

    `template_path` is the chosen persona's `.docx` (optional); both the `.docx`
    and `.pdf` paths borrow its font (the `.pdf` via its `.css` sibling, the
    `.docx` via the same CSS primary family) so the letter matches the résumé
    persona. Unknown formats fall back to `.docx` (the download route validates
    the enum upstream). No JSON Resume sidecar — a cover letter is not a résumé.
    """
    out_dir = Path(base_dir) / username
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    content = _normalize_markdown(content)

    if output_format == ".md":
        path = out_dir / f"cover_letter_{ts}.md"
        path.write_text(content, encoding="utf-8")
    elif output_format == ".pdf":
        path = out_dir / f"cover_letter_{ts}.pdf"
        _render_cover_letter_pdf(content, template_path, path)
    else:
        path = out_dir / f"cover_letter_{ts}.docx"
        _write_cover_letter_docx(content, path, template_path=template_path)

    return str(path)


def _cover_letter_font_name(template_path: str | None) -> str:
    """Resolve the cover letter's body font NAME from the chosen persona.

    Shares one source of truth with the `.pdf` path: `persona_font_family`
    returns the persona CSS's full font stack (e.g. `"Helvetica Neue", Helvetica,
    Arial, sans-serif`); the `.docx` needs a single Word font name, so we take the
    primary family. With no persona / CSS, `persona_font_family` returns the
    neutral business stack, whose primary family is the effective fallback;
    Calibri is only the last-resort guard if that stack were ever empty.
    Deterministic — no LLM.
    """
    from pdf_render import persona_font_family

    css_path = Path(template_path).with_suffix(".css") if template_path else None
    stack = persona_font_family(css_path)
    primary = stack.split(",")[0].strip().strip('"').strip("'")
    return primary or "Calibri"


def _render_cover_letter_pdf(
    content: str,
    template_path: str | None,
    output_pdf_path: Path,
) -> None:
    """Render the cover letter to PDF via the shared business-letter shell.

    Mirrors `_render_pdf_from_json`: resolves the shared
    `personas/cover_letter.html` shell relative to this file, borrows the chosen
    persona's font (its `.css` sibling), and delegates the Chromium pass to
    `pdf_render.render_cover_letter_pdf`. The shell is shared (persona-agnostic),
    so there is no per-persona companion to fall back to.
    """
    from pdf_render import persona_font_family, render_cover_letter_pdf

    repo_root = Path(__file__).resolve().parent
    shell = repo_root / "personas" / "cover_letter.html"
    css_path = Path(template_path).with_suffix(".css") if template_path else None
    font_family = persona_font_family(css_path)
    render_cover_letter_pdf(
        content,
        font_family=font_family,
        template_path=shell,
        output_pdf_path=output_pdf_path,
    )


def _write_cover_letter_docx(
    content: str,
    path: Path,
    template_path: str | None = None,
) -> None:
    """Write the cover letter to a business-letter `.docx`.

    The 2026-05-26 styling decisions (shared with the `.pdf`/preview shell):
      - Persona font (matching the chosen résumé template), plain.
      - Dense, near-single-spaced body — business-letter register, not the
        breathy line-height a résumé uses.
      - TERSER than the résumé: NO centered name banner, NO section styling.
        Every line is a plain Normal paragraph; the date / addressee / salutation
        flow inline as tight stacked paragraphs (no boxed block).

    Distinct from `_write_docx_from_json_resume` (the résumé's structured writer)
    because the cover letter has no name / heading / job-row roles — it is plain
    dense paragraphs. Incidental markdown markers (`#`, bullets) are stripped to
    plain text so a stray heading never renders as a banner. Deterministic — no LLM.
    """
    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name = _cover_letter_font_name(template_path)
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.15  # near single-spaced — business-letter dense
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)  # ≈ the shell's 9pt inter-paragraph gap
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            # Blank lines separate paragraphs; the per-paragraph spacing above
            # carries the rhythm. Emitting an empty paragraph would double gaps.
            continue
        # Strip any incidental heading / bullet marker → plain text (terser intent).
        if stripped.startswith(("# ", "## ", "### ")):
            text = stripped.lstrip("#").strip()
        elif BULLET_RE.match(stripped):
            text = BULLET_RE.sub("", stripped)
        else:
            text = stripped
        p = doc.add_paragraph()
        _add_inline_runs(p, text)  # honors **bold** / *italic*

    doc.save(str(path))


def _add_inline_runs(paragraph: Paragraph, text: str, base_bold: bool = False) -> None:
    """Parse inline **bold** and *italic* markers and add styled runs.

    Handles: ***bold+italic***, **bold**, *italic*, and plain text segments.
    Leaves the paragraph's style intact — only run-level formatting is set.
    """
    segments = _INLINE_RE.split(text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("***") and seg.endswith("***"):
            run = paragraph.add_run(seg[3:-3])
            run.bold = True
            run.italic = True
        elif seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            run = paragraph.add_run(seg[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(seg)
            if base_bold:
                run.bold = True


def _extract_list_numPr(doc: "docx.document.Document") -> CT_NumPr | None:
    """Return a deep copy of the numPr element from the first List Paragraph, or None."""
    for p in doc.paragraphs:
        if p.style and p.style.name == "List Paragraph":
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    return deepcopy(numPr)
    return None


def _clear_body(doc: "docx.document.Document") -> None:
    """Remove all paragraph and table elements from the document body.

    Preserves the final w:sectPr (section/margin properties) which Word
    stores as a direct child of w:body — it is not a w:p or w:tbl element.
    """
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)


def _apply_numPr(paragraph: Paragraph, numPr_template: CT_NumPr) -> None:
    """Attach list numbering properties to a paragraph element."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.insert(0, deepcopy(numPr_template))


# ── Template style capture ──────────────────────────────────────────────────
# Many resumes do not use Word's named heading styles — they apply direct
# formatting (alignment, run sizes, tab stops) to Normal paragraphs. To
# faithfully reproduce these layouts, we walk the template's paragraphs and
# build a dict of role-to-prototype, then apply those prototypes when writing
# the corresponding markdown elements.


def _capture_proto(p: Paragraph) -> dict:
    """Capture a paragraph's formatting into a serializable prototype.

    Captures alignment, vertical spacing, tab stops, and the primary run's
    bold/size. Run italic and color are intentionally not captured — they
    are typically inline (`*text*`) rather than role-based.
    """
    run0 = p.runs[0] if p.runs else None
    pf = p.paragraph_format
    return {
        "alignment": p.alignment,
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "tab_stops": [(t.position.pt, t.alignment) for t in (pf.tab_stops or [])],
        "run_bold": bool(run0.bold) if (run0 and run0.bold is not None) else None,
        "run_size_pt": run0.font.size.pt if (run0 and run0.font.size) else None,
    }


def _capture_template_styles(doc: "docx.document.Document") -> dict:
    """Walk the template's first ~30 paragraphs; classify each by role.

    Heuristics:
    - Centered paragraphs at the very top, in order: name, subtitle, contact.
    - Bold paragraph with a right tab stop: job_title (## job-row).
    - Bold paragraph without tab stops: section_heading.
    - Non-bold paragraph immediately after a job_title: job_subtitle.
    - First non-bold paragraph that isn't header / job_subtitle: body.

    Returns a dict keyed by role; missing roles fall back to the writer's
    built-in defaults so a partial capture still yields a clean output.
    """
    styles: dict = {}
    centered_seen = 0
    last_was_job_title = False

    for p in doc.paragraphs:
        if not p.text.strip():
            continue

        proto = _capture_proto(p)
        is_centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        any_bold = any(r.bold for r in p.runs)
        has_right_tab = any(
            t.alignment == WD_TAB_ALIGNMENT.RIGHT for t in (p.paragraph_format.tab_stops or [])
        )

        # Header zone: centered paragraphs at the top, in order
        if is_centered and centered_seen < 3 and "section_heading" not in styles:
            role = ("name", "subtitle", "contact")[centered_seen]
            styles.setdefault(role, proto)
            centered_seen += 1
            continue

        # Bold + right tab stop = job title with right-aligned date
        if any_bold and has_right_tab:
            styles.setdefault("job_title", proto)
            last_was_job_title = True
            continue

        # Bold without tab stops = section heading
        if any_bold:
            styles.setdefault("section_heading", proto)
            last_was_job_title = False
            continue

        # Non-bold immediately after a job title = job subtitle (smaller font)
        if last_was_job_title:
            styles.setdefault("job_subtitle", proto)
            last_was_job_title = False
            continue

        # Default body paragraph
        styles.setdefault("body", proto)
        last_was_job_title = False

    return styles


def _apply_para_proto(p: Paragraph, proto: dict | None) -> None:
    """Apply paragraph-level formatting from a proto. Skips runs."""
    if not proto:
        return
    if proto.get("alignment") is not None:
        p.alignment = proto["alignment"]
    if proto.get("space_before_pt") is not None:
        p.paragraph_format.space_before = Pt(proto["space_before_pt"])
    if proto.get("space_after_pt") is not None:
        p.paragraph_format.space_after = Pt(proto["space_after_pt"])
    for pos_pt, align in proto.get("tab_stops", []):
        p.paragraph_format.tab_stops.add_tab_stop(Pt(pos_pt), alignment=align)


def _apply_run_proto(run: Run, proto: dict | None) -> None:
    """Apply run-level formatting (bold, font size) from a proto."""
    if not proto:
        return
    if proto.get("run_bold") is not None:
        run.bold = proto["run_bold"]
    if proto.get("run_size_pt"):
        run.font.size = Pt(proto["run_size_pt"])


def _add_inline_runs_with_proto(paragraph: Paragraph, text: str, proto: dict | None) -> None:
    """Like _add_inline_runs but also applies the proto's run-level format to every emitted run.

    The proto format acts as a baseline (inline ** / * still wins over base bold).
    Tab characters in `text` are preserved; the paragraph's tab stops handle them.
    """
    segments = _INLINE_RE.split(text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("***") and seg.endswith("***"):
            run = paragraph.add_run(seg[3:-3])
            run.bold = True
            run.italic = True
            _apply_run_proto(run, {"run_size_pt": (proto or {}).get("run_size_pt")})
        elif seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
            _apply_run_proto(run, {"run_size_pt": (proto or {}).get("run_size_pt")})
        elif seg.startswith("*") and seg.endswith("*"):
            run = paragraph.add_run(seg[1:-1])
            run.italic = True
            _apply_run_proto(run, {"run_size_pt": (proto or {}).get("run_size_pt")})
        else:
            run = paragraph.add_run(seg)
            _apply_run_proto(run, proto)


def _contact_line_items(basics: dict[str, Any]) -> list[str]:
    """Assemble the contact line EXACTLY as `classic.html` does.

    Order: email · phone · profile URLs (scheme-stripped) · website. Keeping
    this identical to the template's header block (personas/bundled/classic.html)
    is what makes the .docx contact line match the on-screen preview.
    """
    items: list[str] = []
    if basics.get("email"):
        items.append(str(basics["email"]))
    if basics.get("phone"):
        items.append(str(basics["phone"]))
    for profile in basics.get("profiles") or []:
        url = str(profile.get("url", ""))
        items.append(url.replace("https://", "").replace("http://", ""))
    if basics.get("url"):
        items.append(str(basics["url"]).replace("https://", "").replace("http://", ""))
    return [i for i in items if i]


def _date_range(start: str | None, end: str | None) -> str:
    """Render a `Start – End` date range (en-dash), tolerating either side missing."""
    start = (start or "").strip()
    end = (end or "").strip()
    if start and end:
        return f"{start} – {end}"
    return start or end


def _entry_header_text(name: str, position: str, start: str | None, end: str | None) -> str:
    r"""Reassemble a `Company, Position\tStart – End` job/degree header line.

    The TAB lets the persona's captured right tab stop right-align the date,
    matching the template's job-header layout; with no template the tab falls
    back to a default stop.
    """
    left = ", ".join(part for part in (name, position) if part)
    dates = _date_range(start, end)
    return f"{left}\t{dates}" if dates else left


def _write_docx_from_json_resume(
    json_doc: dict[str, Any],
    path: Path,
    template_path: str | None = None,
) -> None:
    """Write a JSON Resume document to .docx — the single, non-divergent writer.

    D3 (single source of truth): this consumes the SAME `json_doc`
    (`md_to_json_resume()` output) that the HTML preview and PDF render, and
    walks it in `classic.html`'s section order (header → summary → experience →
    skills → certifications → education → projects). Because the content source
    is the structured document — not a second markdown parse — DOWNLOAD ==
    PREVIEW by construction, and a non-canonically-titled Summary/Skills can no
    longer render in one surface but vanish from the other.

    Persona fidelity is unchanged: when a `.docx` template is supplied, per-role
    typography (alignment, run sizes, tab stops, list numbering) is still
    captured via `_capture_template_styles` and applied through the same
    `_apply_para_proto` / `_add_inline_runs_with_proto` helpers the old writer
    used. Falls back to clean built-in defaults otherwise. Deterministic — no LLM.
    """
    tp = Path(template_path) if template_path else None
    use_template = bool(tp and tp.exists() and tp.suffix.lower() == ".docx")

    if use_template and tp is not None:
        doc = docx.Document(str(tp))
        styles = _capture_template_styles(doc)
        orig_numPr = _extract_list_numPr(doc)
        _clear_body(doc)
    else:
        styles = {}
        orig_numPr = None
        doc = docx.Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

    # ── Per-role emitters (mirror the old writer's per-marker formatting) ──
    def emit_name(text: str) -> None:
        p = doc.add_paragraph()
        if use_template and "name" in styles:
            _apply_para_proto(p, styles["name"])
            _add_inline_runs_with_proto(p, text, styles["name"])
        else:
            _add_inline_runs(p, text, base_bold=True)
            for run in p.runs:
                run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def emit_header_line(text: str, is_contact: bool) -> None:
        p = doc.add_paragraph()
        if use_template:
            role = "contact" if is_contact else "subtitle"
            proto = styles.get(role) or styles.get("contact") or styles.get("subtitle")
            _apply_para_proto(p, proto)
            _add_inline_runs_with_proto(p, text, proto)
        else:
            _add_inline_runs(p, text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def emit_section_heading(text: str) -> None:
        p = doc.add_paragraph()
        if use_template and "section_heading" in styles:
            _apply_para_proto(p, styles["section_heading"])
            _add_inline_runs_with_proto(p, text, styles["section_heading"])
        else:
            _add_inline_runs(p, text, base_bold=True)
            for run in p.runs:
                run.font.size = Pt(13)
            p.paragraph_format.space_after = Pt(2)

    def emit_entry_header(text: str) -> None:
        p = doc.add_paragraph()
        if use_template and "job_title" in styles:
            _apply_para_proto(p, styles["job_title"])
            _add_inline_runs_with_proto(p, text, styles["job_title"])
        else:
            _add_inline_runs(p, text, base_bold=True)
            if not use_template:
                for run in p.runs:
                    run.font.size = Pt(11)

    def emit_entry_summary(text: str) -> None:
        p = doc.add_paragraph()
        if use_template and "job_subtitle" in styles:
            _apply_para_proto(p, styles["job_subtitle"])
            _add_inline_runs_with_proto(p, text, styles["job_subtitle"])
        else:
            proto = styles.get("body") if use_template else None
            if proto:
                _apply_para_proto(p, proto)
                _add_inline_runs_with_proto(p, text, proto)
            else:
                _add_inline_runs(p, text)

    def emit_body(text: str) -> None:
        p = doc.add_paragraph()
        proto = styles.get("body") if use_template else None
        if proto:
            _apply_para_proto(p, proto)
            _add_inline_runs_with_proto(p, text, proto)
        else:
            _add_inline_runs(p, text)

    def emit_bullet(text: str) -> None:
        if use_template and orig_numPr is not None:
            p = doc.add_paragraph(style="List Paragraph")
            _apply_numPr(p, orig_numPr)
        else:
            p = doc.add_paragraph(style="List Bullet")
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        _add_inline_runs(p, text)

    # ── Walk the document in classic.html's section order ─────────────────
    basics: dict[str, Any] = json_doc.get("basics") or {}
    if basics.get("name"):
        emit_name(str(basics["name"]))
    if basics.get("label"):
        emit_header_line(str(basics["label"]), is_contact=False)
    contact_items = _contact_line_items(basics)
    if contact_items:
        emit_header_line(" · ".join(contact_items), is_contact=True)

    summary = basics.get("summary")
    if summary:
        emit_section_heading("Summary")
        for para in str(summary).split("\n\n"):
            para = para.strip()
            if para:
                emit_body(para)

    work = json_doc.get("work") or []
    if work:
        emit_section_heading("Experience")
        for job in work:
            emit_entry_header(
                _entry_header_text(
                    str(job.get("name") or ""),
                    str(job.get("position") or ""),
                    job.get("startDate"),
                    job.get("endDate"),
                )
            )
            if job.get("summary"):
                emit_entry_summary(str(job["summary"]))
            for h in job.get("highlights") or []:
                emit_bullet(str(h))

    skills = json_doc.get("skills") or []
    if skills:
        emit_section_heading("Skills")
        if any(s.get("keywords") for s in skills):
            for s in skills:
                keywords = ", ".join(str(k) for k in (s.get("keywords") or []))
                label = str(s.get("name") or "")
                emit_bullet(f"**{label}:** {keywords}" if keywords else label)
        else:
            line = " · ".join(str(s.get("name") or "") for s in skills if s.get("name"))
            if line:
                emit_body(line)

    certificates = json_doc.get("certificates") or []
    if certificates:
        emit_section_heading("Certifications")
        for c in certificates:
            if c.get("name"):
                emit_bullet(str(c["name"]))

    education = json_doc.get("education") or []
    if education:
        emit_section_heading("Education")
        for ed in education:
            emit_entry_header(
                _entry_header_text(
                    str(ed.get("institution") or ""),
                    str(ed.get("area") or ""),
                    ed.get("startDate"),
                    ed.get("endDate"),
                )
            )
            if ed.get("score"):
                emit_body(str(ed["score"]))

    projects = json_doc.get("projects") or []
    if projects:
        emit_section_heading("Projects")
        for proj in projects:
            if proj.get("name"):
                emit_entry_header(str(proj["name"]))
            if proj.get("description"):
                emit_entry_summary(str(proj["description"]))
            for h in proj.get("highlights") or []:
                emit_bullet(str(h))

    doc.save(str(path))
