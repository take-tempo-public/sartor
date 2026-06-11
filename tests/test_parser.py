"""Unit tests for parser.py — file-format-agnostic resume parsing."""

import pytest

from parser import _infer_sections, parse_resume


class TestInferSections:
    def test_extracts_known_headings(self):
        text = "## Experience\nSomething\n## Education\nMore"
        sections = _infer_sections(text)
        headings = [s["heading"].lower() for s in sections]
        assert any("experience" in h for h in headings)
        assert any("education" in h for h in headings)

    def test_handles_no_headings(self):
        sections = _infer_sections("Just plain text with no resume sections")
        assert isinstance(sections, list)


class TestParseResume:
    def test_unsupported_format_raises(self, tmp_path):
        bogus = tmp_path / "resume.txt"
        bogus.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported format"):
            parse_resume(str(bogus))

    def test_markdown_returns_expected_keys(self, tmp_path):
        md = tmp_path / "resume.md"
        md.write_text("# Jane Smith\n## Experience\nWorked at Acme")
        result = parse_resume(str(md))
        assert set(result.keys()) >= {"text", "format", "sections", "filename", "filepath"}
        assert result["format"] == ".md"
        assert result["filename"] == "resume.md"

    def test_docx_table_cells_are_extracted(self, tmp_path):
        # Table/column-laid-out résumés keep their content in table cells, which
        # the old paragraph-only parser dropped — yielding empty text and a
        # silent zero-experience ingest. The parser must read cell text too.
        import docx
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Acme Corporation"
        table.cell(0, 1).text = "Staff Engineer · 2020-01 to present"
        path = tmp_path / "table_resume.docx"
        doc.save(str(path))

        result = parse_resume(str(path))
        assert "Acme Corporation" in result["text"]
        assert "Staff Engineer" in result["text"]

    def test_docx_interleaves_paragraphs_and_tables_in_order(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Jane Smith")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Built the data platform"
        doc.add_paragraph("References available")
        path = tmp_path / "mixed.docx"
        doc.save(str(path))

        text = parse_resume(str(path))["text"]
        assert text.index("Jane Smith") < text.index("Built the data platform") < text.index("References available")
