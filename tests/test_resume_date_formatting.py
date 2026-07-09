"""Cross-renderer regression guard for the MM-YYYY date presentation fix
(fix/output-identity-and-dates).

Two mechanical bugs this pins down:
  1. Nothing rendered "– Present" for an open-ended role (missing end date).
  2. Raw ISO `YYYY-MM` passed through verbatim into every rendered surface.

`json_resume.format_date_range` is the single canonical presentation-boundary
helper; this suite exercises it through the THREE renderers that must never
disagree — `.docx` (via `generator.generate_resume`), `.md` (same function),
and the live preview HTML (`pdf_render.render_html_string`) — plus the
frozen-composition `.md` path (`json_resume_to_markdown`) — with the SAME
source markdown containing an open-ended role and a closed ISO range.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx
import pytest

from generator import generate_resume
from json_resume import json_resume_to_markdown, md_to_json_resume
from pdf_render import render_html_string

# One closed role (2022-09 -> 2023-05) and one OPEN-ENDED role (2023-06 ->
# nothing) — the exact shape that used to render with no "Present" fallback
# and with raw ISO dates.
RESUME_MD = (
    "# Dana Cole\n"
    "Staff Engineer\n"
    "dana@example.com\n\n"
    "## Summary\n"
    "Platform engineer.\n\n"
    "## Experience\n\n"
    "### Acme, Staff Engineer\t2023-06\n"
    "- Leads the platform team.\n\n"
    "### Beta Corp, Senior Engineer\t2022-09 – 2023-05\n"
    "- Shipped the migration.\n\n"
    "## Education\n\n"
    "### State University, Computer Science\t2016-09 – 2020-05\n"
)

# The bug-pattern guard: a raw ISO YYYY-MM date must never survive into any
# rendered surface.
_ISO_YEAR_MONTH_RE = re.compile(r"\d{4}-\d{2}")


@pytest.fixture
def classic_template_path() -> Path:
    return Path(__file__).resolve().parents[1] / "personas" / "bundled" / "classic.html"


class TestDocxDateFormatting:
    def test_no_raw_iso_dates_and_present_shown(self, tmp_path):
        path = generate_resume(RESUME_MD, ".docx", "dates", base_dir=str(tmp_path))
        text = "\n".join(p.text for p in docx.Document(path).paragraphs if p.text.strip())
        assert not _ISO_YEAR_MONTH_RE.search(text), text
        assert "Present" in text
        # MM-YYYY formatting cases: 2023-06 -> 06-2023, 2022-09 -> 09-2022, etc.
        assert "06-2023" in text
        assert "09-2022 – 05-2023" in text
        assert "09-2016 – 05-2020" in text


class TestMarkdownDateFormatting:
    def test_no_raw_iso_dates_and_present_shown(self, tmp_path):
        path = generate_resume(RESUME_MD, ".md", "dates", base_dir=str(tmp_path))
        text = Path(path).read_text(encoding="utf-8")
        assert not _ISO_YEAR_MONTH_RE.search(text), text
        assert "Present" in text
        assert "06-2023" in text
        assert "09-2022 – 05-2023" in text

    def test_frozen_composition_markdown_serializer_matches(self):
        """json_resume_to_markdown (the frozen-composition .md path) uses the
        SAME canonical helper — no separate, potentially-divergent formatter."""
        doc = md_to_json_resume(RESUME_MD)
        md = json_resume_to_markdown(doc)
        assert not _ISO_YEAR_MONTH_RE.search(md), md
        assert "Present" in md


class TestPreviewHtmlDateFormatting:
    def test_no_raw_iso_dates_and_present_shown(self, classic_template_path):
        doc = md_to_json_resume(RESUME_MD)
        html = render_html_string(doc, html_template_path=classic_template_path)
        assert not _ISO_YEAR_MONTH_RE.search(html), html
        assert "Present" in html
        assert "06-2023" in html
        assert "09-2022" in html and "05-2023" in html


class TestDownloadPreviewDateParity:
    def test_docx_and_preview_agree_on_every_rendered_date(self, tmp_path, classic_template_path):
        doc = md_to_json_resume(RESUME_MD)
        html = render_html_string(doc, html_template_path=classic_template_path)
        path = generate_resume(RESUME_MD, ".docx", "dates", base_dir=str(tmp_path))
        docx_text = "\n".join(p.text for p in docx.Document(path).paragraphs if p.text.strip())
        for token in ("06-2023", "09-2022 – 05-2023", "Present"):
            assert token in html, f"{token!r} missing from preview HTML"
            assert token in docx_text, f"{token!r} missing from .docx"
