"""B2 ATS-conformance structural gate (RELEASE_ARC §Epic B, item 6).

Asserted on GENERATED OUTPUT ``.docx`` (RELEASE_ARC's "output docx" wording),
not only the bundled inputs: single column, no tables, no text boxes or
drawings, no header/footer text, standard section headings only, and every
font name a member of ``json_resume.APPROVED_FONTS``.

The font assertions check allow-list MEMBERSHIP (an exact approved set), not a
deny-list — a new persona preset, emitter, or capture path that introduces an
off-list font fails here without anyone remembering to extend a list
(blast-radius rows 43/45/49/51). ``tests/test_ats_roundtrip.py`` stays the
content-recovery witness on the same outputs; this file is the layout/typography
half.
"""

from __future__ import annotations

from pathlib import Path

import docx as docx_lib
import pytest

from docx_to_persona_html import detect_layout_fidelity
from json_resume import APPROVED_FONTS, map_to_approved_font

REPO_ROOT = Path(__file__).resolve().parent.parent
BUNDLED_DIR = REPO_ROOT / "personas" / "bundled"
BUNDLED_TEMPLATES = sorted(BUNDLED_DIR.glob("*.docx"))

# The writer's own canonical section titles (generator._write_docx_from_json_resume
# walks classic.html's order) — the "standard headings" of the acceptance
# criterion. Anything else showing up as a section heading is a regression.
STANDARD_HEADINGS = {
    "Summary",
    "Experience",
    "Skills",
    "Education",
    "Certifications",
    "Projects",
}

_DOC = {
    "basics": {
        "name": "Casey Rivera",
        "label": "Platform PM",
        "email": "casey@example.com",
        "summary": "A platform PM who ships.",
    },
    "work": [
        {
            "name": "Acme",
            "position": "Senior PM",
            "startDate": "2021-01",
            "endDate": "2023-05",
            "summary": "Owned the platform.",
            "highlights": [
                "Cut churn 20% with the billing rewrite.",
                "Shipped **bold** improvements mid-line.",
            ],
        }
    ],
    "skills": [{"name": "Python"}, {"name": "Kubernetes"}],
    "education": [
        {
            "institution": "State University",
            "area": "Bachelor of Science",
            "studyType": "Computer Science",
            "startDate": "2010-09",
            "endDate": "2014-05",
        }
    ],
    "certificates": [{"name": "PMP"}],
    "projects": [],
}

# One id per template plus the no-template default writer.
_TEMPLATE_PARAMS = [pytest.param(t, id=t.stem) for t in BUNDLED_TEMPLATES] + [
    pytest.param(None, id="no-template")
]


def _generate(tmp_path: Path, template: Path | None):
    from generator import _write_docx_from_json_resume

    out = tmp_path / "out.docx"
    _write_docx_from_json_resume(_DOC, out, template_path=str(template) if template else None)
    return docx_lib.Document(str(out))


class TestGeneratedDocxStructure:
    @pytest.mark.parametrize("template", _TEMPLATE_PARAMS)
    def test_single_column_no_tables_no_textboxes_no_drawings(self, tmp_path, template):
        """Reuses detect_layout_fidelity so there is ONE definition of a clean
        single-column layout (tables / w:cols>1 / txbxContent / drawings)."""
        doc = _generate(tmp_path, template)
        assert detect_layout_fidelity(doc) == "full"

    @pytest.mark.parametrize("template", _TEMPLATE_PARAMS)
    def test_no_header_footer_text(self, tmp_path, template):
        doc = _generate(tmp_path, template)
        for section in doc.sections:
            for container in (section.header, section.footer):
                texts = [p.text.strip() for p in container.paragraphs if p.text.strip()]
                assert texts == [], f"unexpected header/footer text: {texts}"

    @pytest.mark.parametrize("template", _TEMPLATE_PARAMS)
    def test_standard_headings_only(self, tmp_path, template):
        doc = _generate(tmp_path, template)
        texts = {p.text.strip() for p in doc.paragraphs if p.text.strip()}
        emitted_headings = texts & STANDARD_HEADINGS
        # Every section fed in came out under its standard title...
        assert {"Summary", "Experience", "Skills", "Education", "Certifications"} <= (
            emitted_headings
        )
        # ...and no paragraph is a case-variant of a standard heading (the
        # shape a nonstandard-heading regression would take in this writer).
        for text in texts:
            if text.lower() in {h.lower() for h in STANDARD_HEADINGS}:
                assert text in STANDARD_HEADINGS, f"nonstandard heading variant: {text!r}"

    @pytest.mark.parametrize("template", _TEMPLATE_PARAMS)
    def test_every_font_name_on_approved_list(self, tmp_path, template):
        """The exact-set font assertion (blast-radius row 51) + the Normal-style
        witness for row 43: pre-B2, all four bundled templates produced output
        with Normal.font.name unset (probe R2), inheriting Word docDefaults."""
        doc = _generate(tmp_path, template)
        normal_font = doc.styles["Normal"].font.name
        assert normal_font in APPROVED_FONTS, f"Normal style font {normal_font!r} off-list"
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.name is not None:
                    assert run.font.name in APPROVED_FONTS, (
                        f"run font {run.font.name!r} off-list in {p.text[:40]!r}"
                    )

    def test_offlist_template_font_is_mapped_on_apply(self, tmp_path):
        """Rows 43+45 end-to-end: a template whose Normal AND body runs carry an
        off-list font produces output whose fonts are all approved — B1b's font
        carry-through is the path an off-list name would otherwise ride."""
        template = tmp_path / "papyrus.docx"
        src = docx_lib.Document()
        src.styles["Normal"].font.name = "Papyrus"
        heading = src.add_paragraph()
        run = heading.add_run("SECTION")
        run.bold = True
        run.font.name = "Papyrus"
        body = src.add_paragraph()
        body.add_run("Body text in an off-list face.").font.name = "Papyrus"
        src.save(str(template))

        doc = _generate(tmp_path, template)
        assert doc.styles["Normal"].font.name in APPROVED_FONTS
        for p in doc.paragraphs:
            for run in p.runs:
                assert run.font.name is None or run.font.name in APPROVED_FONTS


class TestApprovedFontSources:
    def test_bundled_docx_presets_all_approved(self):
        """Blast-radius row 49: nothing to change in the presets (all approved
        already) — this is the check that STOPS a future preset going off-list."""
        from scripts.build_bundled_templates import PRESETS

        for preset in PRESETS:
            assert preset.font_family in APPROVED_FONTS, (
                f"{preset.filename}: {preset.font_family!r} off-list"
            )

    def test_bundled_css_primaries_all_approved(self):
        """Rows 47/48: every bundled persona CSS leads with an approved family,
        extracted by the SAME persona_font_family the cover-letter .docx and
        .pdf paths read — asserting through the product's own extractor."""
        from pdf_render import persona_font_family

        for css in sorted(BUNDLED_DIR.glob("*.css")):
            stack = persona_font_family(css)
            primary = stack.split(",")[0].strip().strip('"').strip("'")
            assert primary in APPROVED_FONTS, f"{css.name}: primary {primary!r} off-list"

    def test_cover_letter_font_name_maps_offlist_primary(self, tmp_path):
        """Row 46: an uploaded companion CSS predating B2 can still lead
        off-list — the .docx font NAME must map onto the approved list."""
        from generator import _cover_letter_font_name

        docx_path = tmp_path / "persona.docx"
        (tmp_path / "persona.css").write_text(
            "body { font-family: Papyrus, fantasy; }", encoding="utf-8"
        )
        assert _cover_letter_font_name(str(docx_path)) in APPROVED_FONTS


class TestMapToApprovedFont:
    def test_members_pass_through_canonical_casing(self):
        assert map_to_approved_font("arial") == "Arial"
        assert map_to_approved_font("Calibri") == "Calibri"
        assert map_to_approved_font("GEORGIA") == "Georgia"

    def test_known_families_map_to_nearest_neighbor(self):
        assert map_to_approved_font("Helvetica Neue") == "Arial"
        assert map_to_approved_font("Roboto") == "Arial"
        assert map_to_approved_font("Times New Roman") == "Georgia"
        assert map_to_approved_font("Segoe UI") == "Calibri"

    def test_unknown_names_use_serif_heuristic_then_arial(self):
        assert map_to_approved_font("Fancy Serif Display") == "Georgia"
        assert map_to_approved_font("Some Sans-Serif Face") == "Arial"
        assert map_to_approved_font("Comic Sans MS") == "Arial"

    def test_absent_name_is_the_writers_calibri_default(self):
        assert map_to_approved_font(None) == "Calibri"
        assert map_to_approved_font("") == "Calibri"
        assert map_to_approved_font("   ") == "Calibri"

    def test_total_over_approved_list(self):
        for name in ("Papyrus", "Wingdings", "宋体", "12345", "serif"):
            assert map_to_approved_font(name) in APPROVED_FONTS
