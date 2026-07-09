"""Tests for the highest-severity fix in fix/output-identity-and-dates:

  1. `_is_pre_corpus_context` / the /api/generate + /api/generate/stream
     guard — a context missing the corpus-era `application_id` marker is
     rejected with a 409 telling the user to re-analyze, instead of being
     silently replayed.
  2. `_resolve_candidate_identity` + the route's `identity_override` wiring
     — a REAL Candidate DB row's name/email/phone/linkedin/website
     unconditionally overrides whatever the LLM's markdown (or a stale
     context) carried, in the downloaded file AND in the cached
     `last_generated_resume` that feeds the WYSIWYG preview.

Uses the real `generator.generate_resume` (not stubbed) — only the LLM
`generate()` call is stubbed — so the DB round-trip through
`json_resume.apply_identity_override` is exercised end to end.
"""

from __future__ import annotations

import io
import json
from pathlib import Path

import docx
import pytest

import blueprints.generation as bgen


@pytest.fixture
def identity_app(tmp_path, monkeypatch):
    import db.session as db_session

    monkeypatch.setattr(db_session, "DEFAULT_DB_PATH", tmp_path / "identity.sqlite")
    db_session._engine = None
    db_session._SessionLocal = None

    from app import create_app
    from config import Config

    app = create_app(Config(base_dir=tmp_path))
    app.config["TESTING"] = True
    output_dir = tmp_path / "output"
    (tmp_path / "configs" / "alice.config").write_text("{}", encoding="utf-8")
    (output_dir / "alice").mkdir()

    # The LLM's OWN markdown carries a STALE identity — an old website that
    # is no longer on the candidate's DB record (the reported bug's exact
    # mechanism: candidate.online_profile_text / a stale context leaking a
    # website into the header the LLM echoes).
    def _stub_generate(client, context_set, analysis, **kwargs):
        return {
            "resume_content": (
                "# Old Stale Name\n"
                "old-stale@example.com | https://stray-old-site.example\n\n"
                "## Summary\n"
                "Body text.\n"
            ),
            "cover_letter_content": "",
            "changes_made": [],
            "proofread_notes": [],
        }

    monkeypatch.setattr(bgen, "generate", _stub_generate)
    monkeypatch.setattr(bgen, "_get_client", lambda: object())

    from db.models import Application, Base, Candidate
    from db.session import get_engine, get_session

    Base.metadata.create_all(get_engine())
    session = get_session()
    try:
        candidate = Candidate(
            username="alice",
            name="Real Current Name",
            email="real-current@example.com",
            phone="555-0100",
            linkedin_url="https://linkedin.com/in/realcurrent",
            website_url="",  # candidate removed their website — must NOT survive
        )
        session.add(candidate)
        session.flush()
        app_row = Application(
            candidate_id=candidate.id,
            title="Staff Engineer",
            jd_text="...",
            jd_fingerprint="abcd1234",
        )
        session.add(app_row)
        session.commit()
        application_id = app_row.id
    finally:
        session.close()

    ctx_path = output_dir / "alice" / "context_iter0.json"
    ctx_path.write_text(
        json.dumps(
            {
                "resume": {
                    "text": "orig",
                    "filename": "alice.docx",
                    "format": ".docx",
                    "sections": [],
                    "path": "",
                },
                "candidate": {"name": "Ignored — DB is authoritative", "skills": []},
                "job_description": "JD body.",
                "llm_analysis": {"essential_skills": []},
                "deterministic_analysis": {"keyword_overlap": {}},
                "iteration": 0,
                "run_id": "rid",
                "application_id": application_id,
            }
        ),
        encoding="utf-8",
    )
    return app.test_client(), ctx_path, output_dir


class TestIdentityOverrideEndToEnd:
    def test_download_shows_db_identity_not_llm_markdown(self, identity_app):
        client, ctx_path, _output_dir = identity_app
        resp = client.post(
            "/api/generate",
            json={"username": "alice", "context_path": str(ctx_path), "output_format": ".docx"},
        )
        assert resp.status_code == 200, resp.get_json()
        body = resp.get_json()

        doc_text = "\n".join(
            p.text for p in docx.Document(body["resume_path"]).paragraphs if p.text.strip()
        )
        assert "Real Current Name" in doc_text
        assert "real-current@example.com" in doc_text
        assert "Old Stale Name" not in doc_text
        assert "old-stale@example.com" not in doc_text
        assert "stray-old-site.example" not in doc_text

    def test_preview_field_matches_download_not_raw_llm_output(self, identity_app):
        """D3: the resume_preview response field (and last_generated_resume,
        which the WYSIWYG preview route reads back) must show the SAME
        corrected identity as the download — not the raw LLM markdown."""
        client, ctx_path, _output_dir = identity_app
        resp = client.post(
            "/api/generate",
            json={"username": "alice", "context_path": str(ctx_path), "output_format": ".docx"},
        )
        body = resp.get_json()
        assert "Real Current Name" in body["resume_preview"]
        assert "Old Stale Name" not in body["resume_preview"]
        assert "stray-old-site.example" not in body["resume_preview"]

        new_ctx = json.loads(Path(body["context_path"]).read_text(encoding="utf-8"))
        assert "Real Current Name" in new_ctx["last_generated_resume"]
        assert "stray-old-site.example" not in new_ctx["last_generated_resume"]


class TestDownloadEditedIdentityOverride:
    """Walkthrough residuals item 4 — /api/download-edited gains the same
    identity_override wiring run_generation/run_generation_stream already had.

    This route has no context_set in scope (it works from raw edited preview
    content + a bare username), so it resolves the Candidate row by username
    via `_resolve_candidate_identity_by_username` instead of by
    application_id — same identity dict shape, same unconditional override.
    """

    def test_download_edited_shows_db_identity_not_edited_markdown(self, identity_app):
        client, _ctx_path, _output_dir = identity_app
        resp = client.post(
            "/api/download-edited",
            json={
                "username": "alice",
                "content": (
                    "# Hand-Edited Stale Name\n"
                    "hand-edited-stale@example.com | https://stray-old-site.example\n\n"
                    "## Summary\n"
                    "Body text.\n"
                ),
                "type": "resume",
                "original_format": ".docx",
            },
        )
        assert resp.status_code == 200, resp.get_json()
        body = resp.get_json()
        dl = client.get(body["download_url"])
        assert dl.status_code == 200

        doc_text = "\n".join(
            p.text for p in docx.Document(io.BytesIO(dl.data)).paragraphs if p.text.strip()
        )
        assert "Real Current Name" in doc_text
        assert "real-current@example.com" in doc_text
        assert "Hand-Edited Stale Name" not in doc_text
        assert "hand-edited-stale@example.com" not in doc_text
        assert "stray-old-site.example" not in doc_text

    def test_unknown_username_is_a_noop_not_a_failure(self, identity_app):
        """A username with no corpus Candidate row (legacy/file-based) must
        still download — identity_override resolves to None, preserving
        whatever basics the edited markdown itself carries."""
        client, _ctx_path, _output_dir = identity_app
        (Path(client.application.config["CONFIGS_DIR"]) / "ghostwriter.config").write_text(
            "{}", encoding="utf-8"
        )
        (Path(client.application.config["OUTPUT_DIR"]) / "ghostwriter").mkdir(
            parents=True, exist_ok=True
        )
        resp = client.post(
            "/api/download-edited",
            json={
                "username": "ghostwriter",
                "content": "# No Corpus Row\n\n## Summary\nBody text.\n",
                "type": "resume",
                "original_format": ".docx",
            },
        )
        assert resp.status_code == 200, resp.get_json()
        body = resp.get_json()
        dl = client.get(body["download_url"])
        assert dl.status_code == 200
        doc_text = "\n".join(
            p.text for p in docx.Document(io.BytesIO(dl.data)).paragraphs if p.text.strip()
        )
        assert "No Corpus Row" in doc_text


class TestPreCorpusContextGuard:
    def test_generate_rejects_context_missing_application_id(self, identity_app, tmp_path):
        client, ctx_path, output_dir = identity_app
        legacy_ctx = json.loads(ctx_path.read_text(encoding="utf-8"))
        del legacy_ctx["application_id"]
        legacy_path = output_dir / "alice" / "context_legacy.json"
        legacy_path.write_text(json.dumps(legacy_ctx), encoding="utf-8")

        resp = client.post(
            "/api/generate",
            json={"username": "alice", "context_path": str(legacy_path)},
        )
        assert resp.status_code == 409
        body = resp.get_json()
        assert body["needs_reanalyze"] is True

    def test_generate_stream_rejects_context_missing_application_id(self, identity_app, tmp_path):
        client, ctx_path, output_dir = identity_app
        legacy_ctx = json.loads(ctx_path.read_text(encoding="utf-8"))
        del legacy_ctx["application_id"]
        legacy_path = output_dir / "alice" / "context_legacy_stream.json"
        legacy_path.write_text(json.dumps(legacy_ctx), encoding="utf-8")

        resp = client.post(
            "/api/generate/stream",
            json={"username": "alice", "context_path": str(legacy_path)},
        )
        assert resp.status_code == 409

    def test_generate_accepts_context_with_application_id(self, identity_app):
        """Sanity check: the fixture's own context (application_id present)
        is NOT rejected — proves the guard is shape-based, not a blanket 409."""
        client, ctx_path, _output_dir = identity_app
        resp = client.post(
            "/api/generate",
            json={"username": "alice", "context_path": str(ctx_path)},
        )
        assert resp.status_code == 200
