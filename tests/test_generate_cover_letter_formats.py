"""Tests for `generator.generate_cover_letter` format branches
(feat/cover-letter-formats).

The cover letter gained a format param mirroring `generate_resume`: `.md`
(normalized markdown), `.docx` (business-letter writer with the persona font),
and `.pdf` (rendered through the shared `personas/cover_letter.html` shell via
Playwright). These tests are fast + LLM-free + Chromium-free: the `.pdf` *dispatch*
is checked by monkeypatching the renderer; the real Chromium render is covered by
`TestCoverLetterPdfRenderEndToEnd` in `tests/test_pdf_render.py` (slow / gated).
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from generator import generate_cover_letter

REPO_ROOT = Path(__file__).resolve().parents[1]
CLASSIC_DOCX = REPO_ROOT / "personas" / "bundled" / "classic.docx"

SAMPLE_CL = (
    "June 4, 2026\n"
    "Hiring Manager, Acme Corp\n\n"
    "Dear Hiring Manager,\n\n"
    "I rebuilt three distributed systems after scaling a platform to 4M users.\n\n"
    "Sincerely,\nPriya Patel"
)


# -------------------------------------------------------------------
# .md branch
# -------------------------------------------------------------------


class TestCoverLetterMarkdown:
    def test_md_writes_markdown_file(self, tmp_path):
        path = generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".md",
        )
        p = Path(path)
        assert p.suffix == ".md"
        assert p.exists()
        text = p.read_text(encoding="utf-8")
        assert "Dear Hiring Manager," in text
        assert "distributed systems" in text
        assert "Priya Patel" in text

    def test_md_content_is_normalized(self, tmp_path):
        # _normalize_markdown collapses 3+ newlines to 2 and strips trailing space.
        path = generate_cover_letter(
            "Line one.\n\n\n\nLine two.   ",
            "u",
            str(tmp_path),
            output_format=".md",
        )
        text = Path(path).read_text(encoding="utf-8")
        assert "\n\n\n" not in text


# -------------------------------------------------------------------
# .docx branch — business-letter writer
# -------------------------------------------------------------------


class TestCoverLetterDocx:
    def test_default_format_is_docx(self, tmp_path):
        path = generate_cover_letter(SAMPLE_CL, "u", str(tmp_path))
        assert Path(path).suffix == ".docx"
        assert Path(path).exists()

    def test_unknown_format_falls_back_to_docx(self, tmp_path):
        path = generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".rtf",
        )
        assert Path(path).suffix == ".docx"

    def test_docx_contains_body_text(self, tmp_path):
        path = generate_cover_letter(SAMPLE_CL, "u", str(tmp_path), output_format=".docx")
        doc = docx.Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Dear Hiring Manager," in all_text
        assert "distributed systems" in all_text
        assert "Priya Patel" in all_text

    @pytest.mark.skipif(not CLASSIC_DOCX.exists(), reason="bundled classic.docx missing")
    def test_docx_uses_persona_font(self, tmp_path):
        # With a persona template, the Normal font is the persona CSS's primary
        # family — the same source the .pdf uses (classic.css → "Helvetica Neue").
        path = generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".docx",
            template_path=str(CLASSIC_DOCX),
        )
        doc = docx.Document(path)
        assert doc.styles["Normal"].font.name == "Helvetica Neue"

    def test_docx_no_template_uses_default_business_font(self, tmp_path):
        path = generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".docx",
            template_path=None,
        )
        doc = docx.Document(path)
        # No persona CSS → persona_font_family returns the neutral business stack
        # whose primary family is "Helvetica Neue" (Calibri is only the last-resort
        # guard if that stack were ever empty).
        assert doc.styles["Normal"].font.name == "Helvetica Neue"

    def test_docx_is_terser_no_centered_banner(self, tmp_path):
        # Business-letter intent: no centered name banner, no oversized heading runs.
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        path = generate_cover_letter(
            "# Priya Patel\n\nDear Hiring Manager,\n\nBody.",
            "u",
            str(tmp_path),
            output_format=".docx",
        )
        doc = docx.Document(path)
        for p in doc.paragraphs:
            assert p.alignment != WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                assert run.font.size != Pt(16)
        # The incidental "# Priya Patel" heading marker is stripped to plain text.
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "# Priya Patel" not in all_text
        assert "Priya Patel" in all_text


# -------------------------------------------------------------------
# .pdf branch — dispatch only (Chromium-free)
# -------------------------------------------------------------------


class TestCoverLetterPdfDispatch:
    @pytest.mark.skipif(not CLASSIC_DOCX.exists(), reason="bundled classic.docx missing")
    def test_pdf_routes_through_shared_shell_with_persona_font(self, tmp_path, monkeypatch):
        captured: dict = {}

        def _stub(
            cover_letter_markdown,
            *,
            font_family,
            template_path,
            output_pdf_path,
            chromium_args=None,
        ):
            captured["markdown"] = cover_letter_markdown
            captured["font_family"] = font_family
            captured["template_path"] = Path(template_path)
            captured["output_pdf_path"] = Path(output_pdf_path)
            Path(output_pdf_path).write_bytes(b"%PDF-stub")
            return Path(output_pdf_path)

        import pdf_render

        monkeypatch.setattr(pdf_render, "render_cover_letter_pdf", _stub)

        path = generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".pdf",
            template_path=str(CLASSIC_DOCX),
        )

        assert Path(path).suffix == ".pdf"
        # Routed through the shared persona-agnostic business-letter shell.
        assert captured["template_path"].name == "cover_letter.html"
        # Persona font sourced from classic.css (full stack, not just the family).
        assert "Helvetica Neue" in captured["font_family"]
        assert "distributed systems" in captured["markdown"]

    def test_pdf_without_template_uses_default_font(self, tmp_path, monkeypatch):
        captured: dict = {}

        def _stub(
            cover_letter_markdown,
            *,
            font_family,
            template_path,
            output_pdf_path,
            chromium_args=None,
        ):
            captured["font_family"] = font_family
            Path(output_pdf_path).write_bytes(b"%PDF-stub")
            return Path(output_pdf_path)

        import pdf_render
        from pdf_render import _DEFAULT_COVER_LETTER_FONT

        monkeypatch.setattr(pdf_render, "render_cover_letter_pdf", _stub)

        generate_cover_letter(
            SAMPLE_CL,
            "u",
            str(tmp_path),
            output_format=".pdf",
            template_path=None,
        )
        assert captured["font_family"] == _DEFAULT_COVER_LETTER_FONT
