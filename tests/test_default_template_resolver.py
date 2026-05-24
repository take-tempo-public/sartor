"""Tests for `_resolve_default_persona_template_path` (Phase β.1).

The resolver was previously a 3-line hardcode that always returned the
bundled `Classic Single-Column` template. This test suite pins down the
v1.0 priority chain:

  1. Candidate's `is_default = 1` template matching the application's
     `target_role_tag_id`
  2. Candidate's `is_default = 1` template with `primary_role_tag_id IS NULL`
  3. Bundled Classic Single-Column (the original fallback)

Plus the back-compat behavior when no username is passed.
"""

from __future__ import annotations

import pytest


@pytest.fixture
def resolver_app(tmp_path, monkeypatch):
    """Reload app.py against a fresh in-memory DB + temp dirs.

    Mirrors the persona_routes fixture pattern so the migration-seeded
    bundled rows (including Classic Single-Column) are present and the
    Phase β.1 resolver has a real DB to query.
    """
    db_file = tmp_path / "resolver.sqlite"

    import db.session as db_session_mod
    monkeypatch.setattr(db_session_mod, "DEFAULT_DB_PATH", db_file)
    db_session_mod._engine = None
    db_session_mod._SessionLocal = None

    import importlib

    import app as app_module
    importlib.reload(app_module)
    monkeypatch.setattr(app_module, "CONFIGS_DIR", tmp_path / "configs")
    monkeypatch.setattr(app_module, "OUTPUT_DIR", tmp_path / "output")
    monkeypatch.setattr(app_module, "PERSONAS_DIR", tmp_path / "personas")
    monkeypatch.setattr(app_module, "BUNDLED_PERSONAS_DIR", tmp_path / "personas" / "bundled")
    monkeypatch.setattr(app_module, "BASE_DIR", tmp_path)
    (tmp_path / "configs").mkdir()
    (tmp_path / "personas").mkdir()
    (tmp_path / "personas" / "bundled").mkdir()
    (tmp_path / "output").mkdir()

    from db.session import init_db
    init_db(db_file)

    # The seed migration (0002) inserted DB rows pointing at
    # personas/bundled/*.docx, but the actual files don't exist in this
    # tmp_path. _resolve_persona_template_path enforces disk_path.exists()
    # as defense-in-depth, so the fallback to Classic would return None
    # without these stub files. Materialize the minimum needed for tests
    # to verify the resolver returns a real path string.
    from docx import Document
    for filename in ("classic.docx", "modern.docx"):
        target = tmp_path / "personas" / "bundled" / filename
        doc = Document()
        doc.add_paragraph(f"Test stub for bundled {filename}.")
        doc.save(str(target))

    return app_module


def _make_candidate(username="testuser"):
    from db.models import Candidate
    from db.session import get_session
    session = get_session()
    try:
        c = Candidate(username=username, name="Test User")
        session.add(c)
        session.commit()
        return c.id
    finally:
        session.close()


def _make_tag(candidate_id, value, kind="role"):
    from db.models import Tag
    from db.session import get_session
    session = get_session()
    try:
        t = Tag(candidate_id=candidate_id, kind=kind, value=value, display_value=value)
        session.add(t)
        session.commit()
        return t.id
    finally:
        session.close()


def _make_application(candidate_id, role_tag_id=None, title="Senior PM"):
    from db.models import Application
    from db.session import get_session
    session = get_session()
    try:
        a = Application(
            candidate_id=candidate_id,
            title=title,
            jd_text="placeholder JD",
            jd_fingerprint="x" * 16,
            target_role_tag_id=role_tag_id,
        )
        session.add(a)
        session.commit()
        return a.id
    finally:
        session.close()


def _make_user_template(app_module, candidate_id, name, *, is_default=0,
                       primary_role_tag_id=None, filename=None):
    """Drop a minimal valid .docx into personas/{candidate_id}/ + insert row."""
    from docx import Document

    from db.models import PersonaTemplate
    from db.session import get_session

    filename = filename or f"{name.lower().replace(' ', '_')}.docx"
    user_dir = app_module.BASE_DIR / "personas" / str(candidate_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    target = user_dir / filename
    doc = Document()
    doc.add_paragraph(f"User template {name}.")
    doc.save(str(target))

    session = get_session()
    try:
        row = PersonaTemplate(
            candidate_id=candidate_id, name=name,
            path=f"personas/{candidate_id}/{filename}", source="user_upload",
            is_default=is_default, primary_role_tag_id=primary_role_tag_id,
        )
        session.add(row)
        session.commit()
        return row.id, target
    finally:
        session.close()


class TestNoUsername:
    """When called without a username, behavior matches the pre-β.1
    hardcode — always returns Classic Single-Column."""

    def test_returns_classic_when_username_is_none(self, resolver_app):
        result = resolver_app._resolve_default_persona_template_path()
        assert result is not None
        assert "Classic" in result or "classic" in result.lower()

    def test_returns_classic_when_username_is_empty_string(self, resolver_app):
        result = resolver_app._resolve_default_persona_template_path(username="")
        assert result is not None
        assert "Classic" in result or "classic" in result.lower()


class TestNoDefaultsConfigured:
    """A candidate exists but has no `is_default = 1` templates → falls back
    to Classic."""

    def test_known_username_no_defaults(self, resolver_app):
        _make_candidate("alice")
        result = resolver_app._resolve_default_persona_template_path(username="alice")
        assert result is not None
        assert "Classic" in result or "classic" in result.lower()

    def test_unknown_username_falls_through(self, resolver_app):
        result = resolver_app._resolve_default_persona_template_path(username="nobody")
        assert result is not None
        assert "Classic" in result or "classic" in result.lower()


class TestGeneralDefault:
    """Priority 2: candidate has an `is_default = 1` template with
    `primary_role_tag_id IS NULL`. Returned when no role tag matches AND
    when no application_id is supplied."""

    def test_general_default_returned_without_application(self, resolver_app):
        cid = _make_candidate("alice")
        row_id, target_path = _make_user_template(
            resolver_app, cid, "Alice Master", is_default=1,
        )
        result = resolver_app._resolve_default_persona_template_path(username="alice")
        assert result is not None
        assert str(target_path) in result

    def test_general_default_returned_when_application_has_no_role(self, resolver_app):
        cid = _make_candidate("alice")
        app_id = _make_application(cid, role_tag_id=None)
        row_id, target_path = _make_user_template(
            resolver_app, cid, "Alice Master", is_default=1,
        )
        result = resolver_app._resolve_default_persona_template_path(
            username="alice", application_id=app_id,
        )
        assert result is not None
        assert str(target_path) in result


class TestRoleSpecificDefault:
    """Priority 1: candidate has an `is_default = 1` template whose
    `primary_role_tag_id` matches the application's `target_role_tag_id`."""

    def test_role_specific_default_wins_over_general(self, resolver_app):
        cid = _make_candidate("alice")
        design_ic_tag = _make_tag(cid, "design-ic")
        app_id = _make_application(cid, role_tag_id=design_ic_tag)

        # General fallback (would win if role-specific didn't exist)
        _make_user_template(
            resolver_app, cid, "Alice General", is_default=1,
        )
        # Role-specific (should win)
        _, role_target = _make_user_template(
            resolver_app, cid, "Alice Design IC", is_default=1,
            primary_role_tag_id=design_ic_tag, filename="design_ic.docx",
        )

        result = resolver_app._resolve_default_persona_template_path(
            username="alice", application_id=app_id,
        )
        assert result is not None
        assert str(role_target) in result

    def test_falls_to_general_when_role_does_not_match(self, resolver_app):
        cid = _make_candidate("alice")
        design_ic_tag = _make_tag(cid, "design-ic")
        pm_tag = _make_tag(cid, "product-management")  # the application targets this
        app_id = _make_application(cid, role_tag_id=pm_tag)

        # General fallback
        _, general_target = _make_user_template(
            resolver_app, cid, "Alice General", is_default=1,
        )
        # A role-specific default — but for Design IC, not PM
        _make_user_template(
            resolver_app, cid, "Alice Design IC", is_default=1,
            primary_role_tag_id=design_ic_tag, filename="design_ic.docx",
        )

        result = resolver_app._resolve_default_persona_template_path(
            username="alice", application_id=app_id,
        )
        # No PM-specific default → general wins
        assert result is not None
        assert str(general_target) in result


class TestIsolation:
    """A default template for candidate A must NOT be returned for
    candidate B. Tests the candidate_id scoping in the partial-unique
    index works as a query filter too."""

    def test_alice_default_does_not_leak_to_bob(self, resolver_app):
        alice_id = _make_candidate("alice")
        _make_candidate("bob")
        _make_user_template(
            resolver_app, alice_id, "Alice Master", is_default=1,
        )

        # Bob has no templates of his own → must fall back to Classic
        result = resolver_app._resolve_default_persona_template_path(username="bob")
        assert result is not None
        assert "Classic" in result or "classic" in result.lower()
