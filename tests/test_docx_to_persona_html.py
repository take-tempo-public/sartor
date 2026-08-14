"""Tests for docx_to_persona_html — the uploaded-persona HTML+CSS companion generator.

Deterministic module (charter C-6). Covers:
  - round-trip extraction: each bundled .docx re-yields its TypographyPreset knobs
    (clean oracle — the presets ARE the ground truth the .docx was built from);
  - companion emit: .html (skeleton contract preserved, CSS href swapped) + .css
    (uploaded typography) + fidelity sidecar;
  - honest fidelity fallback on non-single-column sources (tables);
  - the integration invariant that makes the preview stop falling back to Classic:
    after generation, `pdf_render.html_template_path_for` resolves the companion.
"""

from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

import docx_to_persona_html as d
from scripts.build_bundled_templates import PRESETS

_REPO_ROOT = Path(__file__).resolve().parents[1]
_BUNDLED = _REPO_ROOT / "personas" / "bundled"


@pytest.mark.parametrize("preset", PRESETS, ids=lambda p: p.filename)
def test_extract_matches_bundled_preset(preset) -> None:
    """Extracting a bundled .docx reconstructs the TypographyPreset it was built from."""
    knobs = d.extract_persona_style(_BUNDLED / preset.filename)

    assert knobs["font_family"] == preset.font_family
    assert knobs["name_pt"] == pytest.approx(preset.name_pt)
    assert knobs["heading_pt"] == pytest.approx(preset.section_heading_pt)
    assert knobs["base_font_pt"] == pytest.approx(preset.body_pt)
    assert knobs["line_spacing"] == pytest.approx(preset.line_spacing)
    assert knobs["heading_uppercase"] is preset.section_heading_uppercase
    assert knobs["heading_small_caps"] is preset.section_heading_small_caps
    assert knobs["heading_underline"] is preset.section_heading_underline
    assert knobs["margin_left_in"] == pytest.approx(preset.margin_inches)
    # All bundled templates are single-column running text.
    assert knobs["layout_fidelity"] == "full"

    # Per-role vertical rhythm (walkthrough residuals, item 1): the builder sets
    # these explicitly on every bundled template, so extraction must recover the
    # exact values, not the CSS defaults.
    assert knobs["header_space_after_pt"] == pytest.approx(preset.section_space_before_pt)
    assert knobs["heading_space_before_pt"] == pytest.approx(preset.section_space_before_pt)
    assert knobs["heading_space_after_pt"] == pytest.approx(preset.section_space_after_pt)
    assert knobs["job_title_space_before_pt"] == pytest.approx(preset.section_space_after_pt)
    # Every bundled template defines a right tab stop for the date column.
    assert knobs["job_title_has_right_tab"] is True


def test_generate_companion_writes_html_css_sidecar(tmp_path) -> None:
    """generate_companion emits .html + .css + fidelity sidecar next to the .docx."""
    docx = tmp_path / "mine.docx"
    shutil.copy(_BUNDLED / "tech.docx", docx)  # Georgia + underlined headings

    result = d.generate_companion(docx)
    assert result is not None
    html_path, css_path = result
    assert html_path == tmp_path / "mine.html"
    assert css_path == tmp_path / "mine.css"

    css = css_path.read_text(encoding="utf-8")
    # The uploaded template's OWN font reaches the preview CSS (not Classic's stack).
    assert "Georgia" in css
    assert "text-decoration: underline" in css  # tech's underlined headings

    # HTML preserves the classic.html Jinja2 contract, swapping only the CSS href.
    html = html_path.read_text(encoding="utf-8")
    assert '<link rel="stylesheet" href="mine.css">' in html
    skeleton = (_BUNDLED / "classic.html").read_text(encoding="utf-8")
    assert html == skeleton.replace('href="classic.css"', 'href="mine.css"')

    sidecar = json.loads((tmp_path / "mine.persona.json").read_text(encoding="utf-8"))
    assert sidecar["layout_fidelity"] == "full"
    # The skeleton stamp rides along so a later skeleton change can invalidate
    # this companion (B1a) — asserted explicitly, not left to widen silently.
    assert sidecar["skeleton_version"] == d.skeleton_version()


def test_generate_companion_resolvable_via_html_template_path_for(tmp_path) -> None:
    """After generation the preview route resolves the companion (no Classic fallback).

    `preview_application_html` / `_render_pdf_from_json` fall back to Classic only
    when `html_template_path_for` returns None. This asserts the fallback no longer
    fires for an uploaded template once its companion exists (walkthrough B2/B3).
    """
    from pdf_render import html_template_path_for

    docx = tmp_path / "owned.docx"
    shutil.copy(_BUNDLED / "spacious.docx", docx)
    assert html_template_path_for(docx) is None  # nothing before generation

    d.generate_companion(docx)
    resolved = html_template_path_for(docx)
    assert resolved == tmp_path / "owned.html"


def test_generate_companion_idempotent(tmp_path) -> None:
    """A second call is a no-op cache hit (companion newer than the .docx)."""
    docx = tmp_path / "mine.docx"
    shutil.copy(_BUNDLED / "modern.docx", docx)

    first = d.generate_companion(docx)
    assert first is not None
    mtime = first[1].stat().st_mtime_ns

    second = d.generate_companion(docx)
    assert second == first
    assert second[1].stat().st_mtime_ns == mtime  # not rewritten


def _add_job_title_paragraph(doc, *, right_tab: bool, space_before_pt: float | None) -> None:
    """Build a bold job-title-shaped paragraph, mirroring a résumé's job-header line.

    `right_tab=False` reproduces what real uploaded résumés routinely do: a bare
    `\\t` before the date with no tab stop defined at all.
    """
    p = doc.add_paragraph()
    if right_tab:
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    run = p.add_run("Acme Corp, Staff Engineer")
    run.bold = True
    p.add_run("\tJanuary 2022 – Present").bold = True


def test_extract_no_explicit_spacing_falls_back_to_css_defaults(tmp_path) -> None:
    """A .docx with no direct paragraph spacing yields None knobs and unchanged CSS.

    Regression guard for the walkthrough-residuals per-role spacing capture: a
    source that never overrides `space_before`/`space_after` (inherits from the
    paragraph style) must render byte-identical rhythm to before the knobs
    existed — `_spacing_css` falls back to the historical px literals.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Casey Rivera").font.size = Pt(18)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("casey@example.com")
    heading = doc.add_paragraph()
    heading.add_run("Experience").bold = True
    _add_job_title_paragraph(doc, right_tab=False, space_before_pt=None)
    docx = tmp_path / "plain.docx"
    doc.save(str(docx))

    knobs = d.extract_persona_style(docx)
    assert knobs["header_space_after_pt"] is None
    assert knobs["heading_space_before_pt"] is None
    assert knobs["heading_space_after_pt"] is None
    assert knobs["job_title_space_before_pt"] is None
    assert knobs["job_title_has_right_tab"] is False

    css = d._build_css(knobs)
    assert "margin-bottom: 20px;" in css  # .resume-header — unchanged default
    assert "margin-top: 20px;" in css  # h2 — unchanged default
    assert "margin-bottom: 8px;" in css  # h2 — unchanged default
    assert "margin-bottom: 14px; page-break-inside" in css  # .job — unchanged default


def test_build_css_honors_captured_spacing(tmp_path) -> None:
    """Explicit space_before/space_after on the source .docx reaches the CSS as pt.

    The job-title line carries a right tab stop here because the role heuristic
    (mirroring `generator._capture_template_styles`) only classifies a line as
    job_title when it is bold AND right-tabbed — without the tab stop the line
    reads as another section heading and its spacing is (correctly) not captured
    as `job_title_space_before_pt`.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Casey Rivera").font.size = Pt(18)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(11)
    p2.add_run("casey@example.com")
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(15)
    heading.paragraph_format.space_after = Pt(5)
    heading.add_run("Experience").bold = True
    _add_job_title_paragraph(doc, right_tab=True, space_before_pt=9)
    docx = tmp_path / "spaced.docx"
    doc.save(str(docx))

    knobs = d.extract_persona_style(docx)
    assert knobs["header_space_after_pt"] == pytest.approx(11)
    assert knobs["heading_space_before_pt"] == pytest.approx(15)
    assert knobs["heading_space_after_pt"] == pytest.approx(5)
    assert knobs["job_title_space_before_pt"] == pytest.approx(9)

    css = d._build_css(knobs)
    assert "margin-bottom: 11pt;" in css  # .resume-header
    assert "margin-top: 15pt;" in css  # h2
    assert "margin-bottom: 5pt;" in css  # h2
    assert "margin-bottom: 9pt; page-break-inside" in css  # .job


def test_build_css_date_layout_matches_docx_reality_no_tab_stop(tmp_path) -> None:
    """No captured right tab stop → the preview must not idealize a right-aligned date.

    Walkthrough residuals item 1 (owner decision): the docx writer stays
    single-tab / no right-alignment work, so a persona whose job-title line has
    no defined tab stop will NOT render its date flush-right in the real .docx
    download. The companion CSS must match that reality instead of the old
    always-`space-between` layout.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Casey Rivera").font.size = Pt(18)
    heading = doc.add_paragraph()
    heading.add_run("Experience").bold = True
    _add_job_title_paragraph(doc, right_tab=False, space_before_pt=None)
    docx = tmp_path / "no_tab.docx"
    doc.save(str(docx))

    knobs = d.extract_persona_style(docx)
    assert knobs["job_title_has_right_tab"] is False
    css = d._build_css(knobs)
    assert "justify-content: flex-start" in css
    assert "justify-content: space-between" not in css


def test_build_css_date_layout_keeps_flush_right_with_tab_stop(tmp_path) -> None:
    """A captured right tab stop keeps the flush-right flex layout (bundled-template case)."""
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Casey Rivera").font.size = Pt(18)
    heading = doc.add_paragraph()
    heading.add_run("Experience").bold = True
    _add_job_title_paragraph(doc, right_tab=True, space_before_pt=None)
    docx = tmp_path / "with_tab.docx"
    doc.save(str(docx))

    knobs = d.extract_persona_style(docx)
    assert knobs["job_title_has_right_tab"] is True
    css = d._build_css(knobs)
    assert "justify-content: space-between" in css


def test_detect_layout_fidelity_flags_tables(tmp_path) -> None:
    """A .docx with a table is not single-column-representable → typography_only."""
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # name role
    run = p.add_run("Casey Rivera")
    run.font.size = Pt(18)
    doc.add_table(rows=2, cols=2)
    docx = tmp_path / "fancy.docx"
    doc.save(str(docx))

    assert d.detect_layout_fidelity(Document(str(docx))) == "typography_only"
    knobs = d.extract_persona_style(docx)
    assert knobs["layout_fidelity"] == "typography_only"
    # Even on a fancy source we still extract typography for a clean single-column
    # preview (matches what _write_docx produces on download).
    result = d.generate_companion(docx)
    assert result is not None
    assert '"layout_fidelity": "typography_only"' in (tmp_path / "fancy.persona.json").read_text(
        encoding="utf-8"
    )


def test_generate_companion_missing_docx_returns_none(tmp_path) -> None:
    """Best-effort: an unreadable .docx yields None (caller falls back to Classic)."""
    assert d.generate_companion(tmp_path / "does-not-exist.docx") is None


# ---------------------------------------------------------------------------
# B1a — stale imported-template companions
#
# A companion is a CLONE of personas/bundled/classic.html, so it goes stale when
# the skeleton changes. The original guard compared only the user's .docx mtime,
# which can never notice OUR file changing: a companion cloned before the
# `date_range` Jinja global landed (2026-07-09, 67b83cc) rendered raw ISO dates
# with no "– Present" forever. Diagnosis:
# docs/dev/diagnosis/b1-stale-template-companions.md
# ---------------------------------------------------------------------------

#: One current role — endDate empty is the DB's NULL-means-current convention,
#: and is exactly the case the missing `date_range` call renders wrongly.
_CURRENT_ROLE_DOC = {
    "basics": {"name": "Ada Lovelace"},
    "work": [
        {
            "name": "Analytical Engines Ltd",
            "position": "Principal Engineer",
            "startDate": "2023-04",
            "endDate": "",
        }
    ],
}


def _stage_pre_date_range_companion(tmp_path: Path, stem: str = "imported") -> Path:
    """Stage a companion exactly as a pre-2026-07-09 import left it, and return the .docx.

    The stale `.html` is the current skeleton with its `date_range(...)` calls
    rewritten back to the raw interpolation the pre-image used, so the fixture
    reproduces the defect's shape without depending on git history being
    reachable (CI checkouts are shallow). The sidecar is byte-for-byte what
    every pre-fix generation wrote: `layout_fidelity` and nothing else.
    """
    docx = tmp_path / f"{stem}.docx"
    shutil.copy(_BUNDLED / "classic.docx", docx)

    skeleton = (_BUNDLED / "classic.html").read_text(encoding="utf-8")
    stale = skeleton.replace(
        "{{ date_range(job.startDate, job.endDate) }}",
        '{{ job.startDate or "" }}{% if job.startDate and job.endDate %} – {% endif %}'
        '{{ job.endDate or "" }}',
    ).replace(
        "{{ date_range(ed.startDate, ed.endDate) }}",
        '{{ ed.startDate or "" }}{% if ed.startDate and ed.endDate %} – {% endif %}'
        '{{ ed.endDate or "" }}',
    )
    assert "date_range(" not in stale, "fixture must reproduce a pre-date_range skeleton"

    docx.with_suffix(".html").write_text(
        stale.replace('href="classic.css"', f'href="{stem}.css"'), encoding="utf-8"
    )
    docx.with_suffix(".css").write_text("/* companion css */\n", encoding="utf-8")
    (tmp_path / f"{stem}.persona.json").write_text(
        json.dumps({"layout_fidelity": "full"}, indent=2), encoding="utf-8"
    )
    return docx


def test_stale_skeleton_companion_is_regenerated(tmp_path) -> None:
    """THE B1a regression: an old-skeleton companion refreshes on next use.

    Fails on HEAD before the fix for two independent reasons — there was no
    resolution entry point that could refresh an existing companion, and the
    mtime guard returned the stale pair unchanged.
    """
    docx = _stage_pre_date_range_companion(tmp_path)
    html_path = docx.with_suffix(".html")
    assert "date_range(" not in html_path.read_text(encoding="utf-8")
    # The guard's old condition is satisfied, so mtime alone would skip forever.
    assert html_path.stat().st_mtime >= docx.stat().st_mtime

    resolved = d.resolve_companion_html(docx)

    assert resolved == html_path
    assert "date_range(" in html_path.read_text(encoding="utf-8")
    sidecar = json.loads((tmp_path / "imported.persona.json").read_text(encoding="utf-8"))
    assert sidecar["skeleton_version"] == d.skeleton_version()
    assert sidecar["layout_fidelity"] == "full"  # the existing marker survives


def test_stale_companion_render_restores_present_and_month_year(tmp_path) -> None:
    """The user-visible acceptance bar: `04/2023 – Present` comes back."""
    from pdf_render import render_html_string

    docx = _stage_pre_date_range_companion(tmp_path)

    before = render_html_string(_CURRENT_ROLE_DOC, html_template_path=docx.with_suffix(".html"))
    assert "Present" not in before
    assert "2023-04" in before  # raw ISO leaks through the stale clone

    resolved = d.resolve_companion_html(docx)
    assert resolved is not None
    after = render_html_string(_CURRENT_ROLE_DOC, html_template_path=resolved)
    assert "04/2023 – Present" in after


def test_generate_companion_regenerates_on_skeleton_stamp_mismatch(tmp_path) -> None:
    """The guard itself keys on the skeleton, not just the user's file timestamps."""
    docx = tmp_path / "mine.docx"
    shutil.copy(_BUNDLED / "modern.docx", docx)
    d.generate_companion(docx)
    html_path = docx.with_suffix(".html")
    sidecar = tmp_path / "mine.persona.json"

    # Freshly generated: a second call is still a no-op cache hit.
    mtime = html_path.stat().st_mtime_ns
    d.generate_companion(docx)
    assert html_path.stat().st_mtime_ns == mtime

    # Now pretend it was cloned from a different skeleton.
    sidecar.write_text(
        json.dumps({"layout_fidelity": "full", "skeleton_version": "0" * 16}, indent=2),
        encoding="utf-8",
    )
    result = d.generate_companion(docx)

    assert result is not None
    assert json.loads(sidecar.read_text(encoding="utf-8"))["skeleton_version"] == (
        d.skeleton_version()
    )


@pytest.mark.parametrize(
    "sidecar_body",
    [
        pytest.param("{not valid json", id="truncated"),
        pytest.param("[]", id="non-dict"),
    ],
)
def test_unreadable_sidecar_is_treated_as_stale_not_bundled(tmp_path, sidecar_body) -> None:
    """A present-but-unreadable sidecar is ours (presence is the ownership test) —

    not a bundled hand-authored companion — so it must self-heal like any other
    stale companion, not freeze forever the way an absent sidecar (genuinely
    bundled) correctly does.
    """
    docx = _stage_pre_date_range_companion(tmp_path)
    html_path = docx.with_suffix(".html")
    sidecar_path = tmp_path / "imported.persona.json"
    sidecar_path.write_text(sidecar_body, encoding="utf-8")

    resolved = d.resolve_companion_html(docx)

    assert resolved == html_path
    assert "date_range(" in html_path.read_text(encoding="utf-8")
    sidecar = json.loads(sidecar_path.read_text(encoding="utf-8"))
    assert sidecar["skeleton_version"] == d.skeleton_version()


def test_current_companion_not_regenerated(tmp_path) -> None:
    """The other half of the acceptance bar: no needless .docx re-parse per preview.

    A stamp bug that invalidated every companion would turn each preview request
    into a python-docx parse — this asserts the resolve path leaves a current
    companion byte-identical.
    """
    docx = tmp_path / "mine.docx"
    shutil.copy(_BUNDLED / "tech.docx", docx)
    generated = d.generate_companion(docx)
    assert generated is not None
    html_path, css_path = generated
    html_mtime = html_path.stat().st_mtime_ns
    css_mtime = css_path.stat().st_mtime_ns

    assert d.resolve_companion_html(docx) == html_path
    assert d.resolve_companion_html(docx) == html_path

    assert html_path.stat().st_mtime_ns == html_mtime
    assert css_path.stat().st_mtime_ns == css_mtime


def test_resolve_companion_html_does_not_touch_bundled_companions() -> None:
    """Hand-authored bundled companions are NOT classic-skeleton clones — never rewrite them.

    They carry no `.persona.json`, and that absence is the ownership test the
    refresh keys on. Without it, adding regenerate-on-mismatch would overwrite
    modern/tech/spacious with the Classic skeleton the moment a preview resolved
    a bundled persona.
    """
    for stem in ("classic", "modern", "spacious", "tech"):
        assert not (_BUNDLED / f"{stem}.persona.json").exists(), (
            "a bundled sidecar would make the ownership test unsound"
        )

    before = {
        stem: (_BUNDLED / f"{stem}.html").read_bytes() for stem in ("modern", "spacious", "tech")
    }

    for stem in ("modern", "spacious", "tech"):
        resolved = d.resolve_companion_html(_BUNDLED / f"{stem}.docx")
        assert resolved == _BUNDLED / f"{stem}.html"

    for stem, content in before.items():
        assert (_BUNDLED / f"{stem}.html").read_bytes() == content


def test_resolve_companion_html_generates_when_absent(tmp_path) -> None:
    """Unchanged behavior for the case the old resolve-then-generate pair handled."""
    docx = tmp_path / "fresh.docx"
    shutil.copy(_BUNDLED / "spacious.docx", docx)
    assert not docx.with_suffix(".html").exists()

    resolved = d.resolve_companion_html(docx)

    assert resolved is not None
    assert resolved == docx.with_suffix(".html")
    assert resolved.exists()


def test_resolve_companion_html_missing_docx_returns_none(tmp_path) -> None:
    """No companion and no readable .docx → None, so callers fall back to Classic."""
    assert d.resolve_companion_html(tmp_path / "does-not-exist.docx") is None


def test_resolve_companion_html_keeps_stale_companion_when_refresh_fails(tmp_path) -> None:
    """A failed refresh must not downgrade the preview to bundled Classic.

    Returning None here would swap the user's own typography for Classic's on a
    transient failure; returning the stale companion keeps today's behavior.
    """
    docx = _stage_pre_date_range_companion(tmp_path)
    html_path = docx.with_suffix(".html")
    stale_bytes = html_path.read_bytes()
    docx.unlink()  # regeneration cannot read the source .docx

    assert d.resolve_companion_html(docx) == html_path
    assert html_path.read_bytes() == stale_bytes


def test_skeleton_version_tracks_the_shipped_skeleton() -> None:
    """The stamp is a content hash, so it cannot drift from the skeleton it describes."""
    import hashlib

    expected = hashlib.sha256((_BUNDLED / "classic.html").read_bytes()).hexdigest()[:16]
    assert d.skeleton_version() == expected
