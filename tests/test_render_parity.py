"""Branch 1 — single-render-engine parity + section-title fidelity (D3).

Two invariants this suite pins:

1. **Download == preview.** The `.docx` download renders from the SAME
   `md_to_json_resume()` document the HTML/PDF preview renders — not a second,
   divergent markdown parse. We prove it by comparing the JSON Resume sidecar
   `generate_resume()` writes (the download's source of truth) against
   `md_to_json_resume()` of the normalized content (the preview's source of
   truth), and by confirming every preview bullet/summary/skill appears in the
   generated `.docx`.

2. **Non-canonical section titles survive.** A résumé that titles its sections
   "Professional Summary" / "Core Competencies" / "Professional Experience"
   (very common, and exactly what plain Word imports produce) must land in the
   canonical JSON Resume fields, not silently drop to `meta.sartor.unparsed`.

3. **Template typography reaches the download** (`fix/b1-education-render`).
   Most real résumés carry their typeface as direct run formatting rather than a
   named style; `_capture_proto` has to read it or the download comes back in the
   theme default. Lives here because it is the same "the download is faithful to
   its source" invariant as (1), measured on the generated `.docx` rather than
   against the preview.
"""

from __future__ import annotations

import json
from pathlib import Path

import docx
import pytest

from generator import _normalize_markdown, generate_resume
from json_resume import _SECTION_MAP, md_to_json_resume

# Uses non-canonical (but ubiquitous) headings + metric-bearing bullets — the
# shape a plain Word résumé import produces.
RESUME_MD = (
    "# Robert Cooksey\n"
    "Product Manager\n"
    "robert@example.com | 555-010-2200 | linkedin.com/in/robert\n\n"
    "## Professional Summary\n"
    "Product leader with 10+ years shipping hardware + software.\n\n"
    "## Professional Experience\n\n"
    "### Core Impact, Co-Founder & Head of Product\t2021 – Present\n"
    "- Owned product vision and success metrics (engagement, retention, NPS).\n"
    "- Grew pipeline via a sales flywheel producing $1M in adoption.\n\n"
    "### Intel, Product Experience Architect\t2016 – 2021\n"
    "- Compressed a 3-4 week documentation cycle to sprint-level integration.\n\n"
    "## Core Competencies\n"
    "Product Strategy · Roadmapping · PRDs\n"
)


class TestSectionTitleAliases:
    @pytest.mark.parametrize(
        ("heading", "canonical"),
        [
            ("Professional Summary", "_summary"),
            ("Summary of Qualifications", "_summary"),
            ("Profile", "_summary"),
            ("Professional Experience", "work"),
            ("Work History", "work"),
            ("Technical Skills", "skills"),
            ("Core Competencies", "skills"),
            ("Areas of Expertise", "skills"),
        ],
    )
    def test_alias_maps_to_canonical_key(self, heading: str, canonical: str) -> None:
        assert _SECTION_MAP[heading.lower()] == canonical

    def test_noncanonical_summary_and_skills_are_captured(self) -> None:
        jr = md_to_json_resume(RESUME_MD)
        assert jr["basics"]["summary"].startswith("Product leader")
        assert [s["name"] for s in jr["skills"]] == ["Product Strategy", "Roadmapping", "PRDs"]
        assert len(jr["work"]) == 2
        # Nothing was silently dropped to the unparsed bucket.
        assert jr["meta"]["sartor"]["unparsed"] == []


class TestDownloadPreviewParity:
    def test_docx_sidecar_equals_preview_json(self, tmp_path: Path) -> None:
        """The JSON Resume sidecar (download's source) equals the preview's source."""
        path = generate_resume(RESUME_MD, ".docx", "parity", base_dir=str(tmp_path))
        sidecar = Path(path).with_suffix(".jsonresume.json")
        got = json.loads(sidecar.read_text(encoding="utf-8"))
        expected = md_to_json_resume(_normalize_markdown(RESUME_MD))
        assert got == expected

    def test_docx_carries_every_preview_bullet_and_field(self, tmp_path: Path) -> None:
        jr = md_to_json_resume(_normalize_markdown(RESUME_MD))
        path = generate_resume(RESUME_MD, ".docx", "parity", base_dir=str(tmp_path))
        paras = [p.text for p in docx.Document(path).paragraphs if p.text.strip()]
        text = "\n".join(paras)

        # Every preview bullet made it into the download (metrics included).
        for job in jr["work"]:
            for highlight in job.get("highlights", []):
                assert highlight in text, f"bullet missing from .docx: {highlight}"
        assert jr["basics"]["summary"] in text
        assert "Product Strategy" in text

    def test_docx_emits_canonical_headings_not_source_titles(self, tmp_path: Path) -> None:
        """The writer renders canonical section names, so preview and download agree."""
        path = generate_resume(RESUME_MD, ".docx", "parity", base_dir=str(tmp_path))
        paras = [p.text for p in docx.Document(path).paragraphs if p.text.strip()]
        assert "Summary" in paras and "Professional Summary" not in paras
        assert "Experience" in paras and "Professional Experience" not in paras
        assert "Skills" in paras and "Core Competencies" not in paras


class TestSectionSpacing:
    """O1a (round-2 quick win): the .docx writer inserts a blank-paragraph
    spacer between top-level sections and between consecutive work entries, so
    the output stops reading as a dense wall of text. On the default (no
    template) path, no role carries captured spacing, so every spacer fires —
    that is the invariant pinned here. Content parity is unaffected (the
    existing parity tests filter `p.text.strip()`, so spacers are invisible to
    them and every bullet/field still lands)."""

    @staticmethod
    def _all_paras(path: str) -> list[str]:
        # Include EMPTY paragraphs — the spacers are empty, so we must not
        # filter them out here (unlike the content-parity tests).
        return [p.text for p in docx.Document(path).paragraphs]

    def test_blank_spacer_before_each_later_section(self, tmp_path: Path) -> None:
        path = generate_resume(RESUME_MD, ".docx", "spacing", base_dir=str(tmp_path))
        paras = self._all_paras(path)
        # The first section (Summary) gets no leading spacer; every later
        # section heading is immediately preceded by a blank paragraph.
        for heading in ("Experience", "Skills"):
            assert heading in paras, f"missing section heading: {heading}"
            idx = paras.index(heading)
            assert idx > 0 and paras[idx - 1] == "", (
                f"expected a blank spacer before the '{heading}' heading"
            )
        # Summary is first — it must NOT be preceded by a spacer.
        s_idx = paras.index("Summary")
        assert paras[s_idx - 1] != "", "first section (Summary) should have no leading spacer"

    def test_blank_spacer_between_work_entries(self, tmp_path: Path) -> None:
        path = generate_resume(RESUME_MD, ".docx", "spacing", base_dir=str(tmp_path))
        paras = self._all_paras(path)
        # RESUME_MD has two work entries; a blank paragraph separates them.
        first = next(i for i, t in enumerate(paras) if t.startswith("Core Impact"))
        second = next(i for i, t in enumerate(paras) if t.startswith("Intel"))
        assert first < second
        assert "" in paras[first + 1 : second], (
            "expected a blank spacer between the two work entries"
        )

    def test_spacers_do_not_disturb_content_parity(self, tmp_path: Path) -> None:
        """Every non-blank line still matches the preview source (spacers add
        emptiness, never text)."""
        jr = md_to_json_resume(_normalize_markdown(RESUME_MD))
        path = generate_resume(RESUME_MD, ".docx", "spacing", base_dir=str(tmp_path))
        text = "\n".join(p.text for p in docx.Document(path).paragraphs if p.text.strip())
        for job in jr["work"]:
            for highlight in job.get("highlights", []):
                assert highlight in text
        assert jr["basics"]["summary"] in text


class TestAtsScrubAndIdentityOverrideParity:
    """fix/output-identity-and-dates: the ATS scrub and identity override run
    inside generate_resume() itself (right after md_to_json_resume), so
    .docx / .md / the jsonresume.json sidecar can never disagree."""

    UNSAFE_MD = (
        "# Dana [QA] Cole\n"
        'Staff Engineer "the closer"\n'
        "dana@example.com\n\n"
        "## Summary\n"
        "Shipped {v2} with <b>bold</b> claims and <50ms latency.\n\n"
        "## Experience\n\n"
        "### Acme, Staff Engineer\t2022-01 – present\n"
        "- Cut p99 latency to <50ms using C++ and C#.\n"
    )

    def test_docx_md_and_sidecar_agree_on_scrubbed_text(self, tmp_path: Path) -> None:
        docx_path = generate_resume(self.UNSAFE_MD, ".docx", "scrub", base_dir=str(tmp_path))
        md_path = generate_resume(self.UNSAFE_MD, ".md", "scrub", base_dir=str(tmp_path))
        sidecar = Path(docx_path).with_suffix(".jsonresume.json")
        sidecar_doc = json.loads(sidecar.read_text(encoding="utf-8"))

        docx_text = "\n".join(p.text for p in docx.Document(docx_path).paragraphs if p.text.strip())
        md_text = Path(md_path).read_text(encoding="utf-8")

        for surface_name, text in (("docx", docx_text), ("md", md_text)):
            assert "[" not in text and "]" not in text, surface_name
            assert "{" not in text and "}" not in text, surface_name
            assert '"' not in text, surface_name
            assert "<b>" not in text, surface_name
            # tag-shaped <...> stripped, but a bare "<50ms" (no closing '>')
            # and C++/C# (neither char is in the strip set) survive.
            assert "<50ms" in text, surface_name
            assert "C++" in text and "C#" in text, surface_name

        assert sidecar_doc["meta"]["sartor"]["ats_scrubbed"]

    def test_identity_override_applies_to_docx_and_md_alike(self, tmp_path: Path) -> None:
        identity = {
            "name": "Real Name",
            "email": "real@example.com",
            "phone": "",
            "linkedin_url": "",
            "website_url": "",
        }
        stale_md = (
            "# Old Name\nold@example.com | https://stray-site.example\n\n## Summary\nBody text.\n"
        )
        docx_path = generate_resume(
            stale_md, ".docx", "identity", base_dir=str(tmp_path), identity_override=identity
        )
        md_path = generate_resume(
            stale_md, ".md", "identity", base_dir=str(tmp_path), identity_override=identity
        )
        docx_text = "\n".join(p.text for p in docx.Document(docx_path).paragraphs if p.text.strip())
        md_text = Path(md_path).read_text(encoding="utf-8")
        for surface_name, text in (("docx", docx_text), ("md", md_text)):
            assert "Real Name" in text, surface_name
            assert "real@example.com" in text, surface_name
            assert "Old Name" not in text, surface_name
            assert "old@example.com" not in text, surface_name
            assert "stray-site.example" not in text, surface_name


# ---------------------------------------------------------------------
# Education: studyType reaches BOTH surfaces (fix/b1-education-render)
# ---------------------------------------------------------------------

CLASSIC_HTML = Path(__file__).resolve().parents[1] / "personas" / "bundled" / "classic.html"

# `Institution, Area — StudyType` — the em dash (U+2014) is the area/studyType
# separator; the en dash (U+2013) between the dates is a different character on
# purpose (json_resume.EDUCATION_FIELD_SEPARATOR).
EDUCATION_MD = (
    "# Jane Doe\n"
    "\n"
    "## Education\n"
    "### State University, Bachelor of Science — Computer Science\t2010 – 2014\n"
)


class TestEducationRenderParity:
    """Education was outside this file's coverage, which is how four render surfaces
    (Classic, Spacious, the `.docx` writer, and the markdown round-trip) drifted away
    from Modern and Tech unnoticed
    (`docs/dev/diagnosis/b1-education-render.md`). Covering it here means a future
    change to one education surface and not the others fails loudly.
    """

    def test_parse_recovers_both_education_fields(self) -> None:
        ed = md_to_json_resume(EDUCATION_MD)["education"][0]
        assert ed == {
            "institution": "State University",
            "area": "Bachelor of Science",
            "studyType": "Computer Science",
            "startDate": "2010",
            "endDate": "2014",
        }

    def test_docx_and_preview_agree_on_education(self, tmp_path: Path) -> None:
        """Download == preview, for education specifically."""
        from pdf_render import render_html_string

        path = generate_resume(EDUCATION_MD, ".docx", "edu", base_dir=str(tmp_path))
        sidecar = json.loads(Path(path).with_suffix(".jsonresume.json").read_text(encoding="utf-8"))
        docx_text = "\n".join(p.text for p in docx.Document(path).paragraphs if p.text.strip())
        html = render_html_string(sidecar, html_template_path=CLASSIC_HTML)

        assert "State University, Bachelor of Science — Computer Science" in docx_text
        for surface_name, text in (("docx", docx_text), ("html", html)):
            assert "Bachelor of Science" in text, surface_name
            assert "Computer Science" in text, surface_name
            # Render both, never flip: area leads, studyType follows.
            assert text.index("Bachelor of Science") < text.index("Computer Science")

    def test_md_download_matches_the_docx_header(self, tmp_path: Path) -> None:
        md_path = generate_resume(EDUCATION_MD, ".md", "edu", base_dir=str(tmp_path))
        md_text = Path(md_path).read_text(encoding="utf-8")
        assert "### State University, Bachelor of Science — Computer Science\t" in md_text


# ---------------------------------------------------------------------
# Template typeface capture (fix/b1-education-render)
# ---------------------------------------------------------------------


def _typography_template(path: Path, font: str | None) -> Path:
    """Write a .docx whose typography is DIRECT run formatting, not named styles.

    Paragraph order matches what `generator._capture_template_styles` classifies:
    three centered (name / subtitle / contact), a bold heading, a bold
    right-tab-stopped job title, then two plain lines (job_subtitle, body).
    `font=None` writes the same shape with no direct typeface at all.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Pt

    doc = docx.Document()

    def para(text: str, *, size: float, centered=False, bold=False, right_tab=False):
        p = doc.add_paragraph()
        if centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if right_tab:
            p.paragraph_format.tab_stops.add_tab_stop(Pt(468), alignment=WD_TAB_ALIGNMENT.RIGHT)
        run = p.add_run(text)
        run.bold = bold or None
        run.font.size = Pt(size)
        if font:
            run.font.name = font

    para("NAME", size=18, centered=True)
    para("Subtitle", size=12, centered=True)
    para("a@b.com", size=10, centered=True)
    para("EXPERIENCE", size=13, bold=True)
    para("Acme, Engineer\t2015 - 2020", size=11, bold=True, right_tab=True)
    para("A role subtitle", size=11)
    para("A body paragraph", size=11)

    doc.save(str(path))
    return path


EMPHASIS_MD = (
    "# Jane Doe\n"
    "\n"
    "## Experience\n"
    "### Acme, Engineer\t2015 - 2020\n"
    "Led **large** cross-functional teams.\n"
    "- Shipped a thing.\n"
)


class TestTemplateFontCapture:
    """A template's direct-run typeface must reach the download.

    `_capture_proto` read bold and size and stopped, so a résumé whose typography
    is direct formatting came back in the theme default
    (`docs/dev/diagnosis/b1-education-render.md` O-6). The HTML-companion path
    has read `run.font.name` since it shipped (`docx_to_persona_html.py:211`) —
    these pin the docx path to the same behavior.
    """

    @staticmethod
    def _runs_by_text(path: str) -> dict[str, list[str | None]]:
        return {
            p.text: [r.font.name for r in p.runs]
            for p in docx.Document(path).paragraphs
            if p.text.strip()
        }

    def test_capture_records_the_direct_run_font(self, tmp_path: Path) -> None:
        from generator import _capture_template_styles

        tpl = _typography_template(tmp_path / "georgia.docx", "Georgia")
        styles = _capture_template_styles(docx.Document(str(tpl)))
        assert {"name", "subtitle", "contact", "section_heading", "job_title"} <= set(styles)
        for role, proto in styles.items():
            assert proto["run_font_name"] == "Georgia", role

    def test_font_reaches_the_generated_download(self, tmp_path: Path) -> None:
        tpl = _typography_template(tmp_path / "georgia.docx", "Georgia")
        path = generate_resume(
            EMPHASIS_MD, ".docx", "font", base_dir=str(tmp_path), template_path=str(tpl)
        )
        runs = self._runs_by_text(path)
        assert runs["Jane Doe"] == ["Georgia"]
        assert runs["Experience"] == ["Georgia"]
        assert any(text.startswith("Acme, Engineer") for text in runs)

    def test_font_survives_an_inline_emphasis_boundary(self, tmp_path: Path) -> None:
        """A typeface must not drop mid-line at a `**bold**` segment."""
        tpl = _typography_template(tmp_path / "georgia.docx", "Georgia")
        path = generate_resume(
            EMPHASIS_MD, ".docx", "font", base_dir=str(tmp_path), template_path=str(tpl)
        )
        runs = self._runs_by_text(path)
        summary = next(fonts for text, fonts in runs.items() if text.startswith("Led "))
        assert len(summary) > 1, "expected the summary to split into multiple runs"
        assert set(summary) == {"Georgia"}

    def test_template_without_a_direct_font_invents_none(self, tmp_path: Path) -> None:
        """Capture must not fabricate a typeface — inheriting from a style is valid."""
        from generator import _capture_template_styles

        tpl = _typography_template(tmp_path / "plain.docx", None)
        styles = _capture_template_styles(docx.Document(str(tpl)))
        assert all(proto["run_font_name"] is None for proto in styles.values())

        path = generate_resume(
            EMPHASIS_MD, ".docx", "font", base_dir=str(tmp_path), template_path=str(tpl)
        )
        assert all(font is None for fonts in self._runs_by_text(path).values() for font in fonts)
