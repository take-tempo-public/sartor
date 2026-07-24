"""Drive the wizard to produce the 10 manifest screenshots.

Generates a synthetic Priya master .docx (matching the worked
example at docs/walkthrough_example.md), walks through all six
wizard steps via Playwright, and captures the 10 PNGs per
docs/ux/screenshot_capture.md.

Cost: ~$0.27 in Anthropic API spend per run (full wizard pass:
extract_experiences + analyze + clarify + recommend_bullets +
recommend_summaries + critique × N + generate + cover letter).
Time: ~6-8 minutes including LLM waits.

Prereqs: app must be running (`python app.py` in another shell),
.api_key in place, Chromium installed (`python -m playwright
install chromium`).

Usage:
    python -m scripts.capture_screenshots
    python -m scripts.capture_screenshots --headless
    python -m scripts.capture_screenshots --keep-user
    python -m scripts.capture_screenshots --headless --smoke

Cleanup: by default the script deletes the demo user and its
on-disk artefacts after capture. Pass --keep-user to retain them
(useful when iterating on capture states).

--smoke: stop after Step 1 (setup + Job/Analyze) instead of running
the full 10-shot pass. This is the periodic drift check driven by
.github/workflows/capture-smoke.yml (monthly cron + manual
workflow_dispatch) — it exists because this script had zero
automated coverage and silently accumulated three independent
staleness bugs over ~7 weeks before anyone noticed (see
docs/dev/diagnosis/capture-screenshots-welcome-modal.md). Cost is
one real Haiku corpus-extract + one real Sonnet analyze call.
"""

from __future__ import annotations

import argparse
import contextlib
import shutil
import sys
import time
from pathlib import Path

from docx import Document
from playwright.sync_api import Page, sync_playwright

from ui_pages import (
    CorpusPage,
    UserPickerPage,
    WizardClarifyPage,
    WizardComposePage,
    WizardGeneratePage,
    WizardJobPage,
    WizardOutputPage,
    WizardTemplatePage,
)
from ui_pages import selectors as S

REPO = Path(__file__).resolve().parent.parent
SHOTS = REPO / "docs" / "screenshots"
TMP = REPO / "scripts" / "_tmp_capture"
APP_URL = "http://localhost:5000"
DEMO_USER = "demo"
VIEWPORT = {"width": 1440, "height": 900}

# Generous timeouts — analyze and generate both routinely take 30-60s.
LLM_TIMEOUT_MS = 120_000
SHORT_TIMEOUT_MS = 15_000


PRIYA_JD = """\
Senior Backend Engineer, Platform — Vertica Logistics

We're hiring a senior IC to own the platform that powers our
real-time shipment tracking and routing engine. You'll work in
Python on a Postgres + AWS stack, with heavy use of Kafka for
our event backbone. Comfort with Kafka topic design, partition
strategy, and consumer-lag debugging is essential.

You'll partner with our SRE and data teams, lead architectural
reviews, and mentor a team of 6 mid- and junior-level engineers
on the platform group. We need someone who has led migrations
of similar scope — moving production systems off legacy queues
onto Kafka, refactoring monolithic services into event-driven
ones. Kafka experience is non-negotiable; we use Kafka across
12 topics serving 8k events/sec at peak.

Bonus: Kubernetes experience, prior work on logistics or
supply-chain platforms.

Stack: Python 3.11+, Postgres 14, AWS (ECS, RDS, S3), Kafka,
Datadog, Terraform. We deploy via GitHub Actions. The team
runs an on-call rotation but expects engineers to invest in
reducing pages rather than absorbing them.

Compensation is competitive for a senior IC role in our region.
We're remote-first; HQ is in Seattle if you want to come in.
"""

SAMPLE_REFINE_NOTE = "emphasize the team-lead role more in the summary"


def write_priya_docx(path: Path) -> None:
    """Synthetic master résumé matching docs/walkthrough_example.md.

    Generates a minimal but parser-friendly .docx with three
    experiences, ~8 bullets each, one Kafka-passing-mention bullet
    on the Helix card, no team-size claim anywhere. Also drops a
    copy into resumes/<DEMO_USER>/ so corpus_import.run_import
    finds it during the corpus seed.
    """
    doc = Document()
    doc.add_heading("Priya Sharma", level=0)
    doc.add_paragraph("priya.sharma@example.com · linkedin.com/in/priya-sharma")

    doc.add_heading("Experience", level=1)

    doc.add_heading("Helix Logistics — Senior Backend Engineer", level=2)
    doc.add_paragraph("2023 – Present")
    for bullet in [
        "Built and maintained Python services on the order-routing platform serving 4M+ daily shipments.",
        "Owned the order-events pipeline; helped migrate it off AWS SQS to Kafka over a six-month dual-write/cutover phase.",
        "Designed Postgres schemas and partition strategies for the routing engine; reduced p99 query latency from 1.2s to 180ms.",
        "Led architecture reviews for the platform team; wrote four ADRs that became the team's default patterns.",
        "Mentored junior and mid-level engineers through pairing, code review, and brown-bag tech talks.",
        "Built incident-response runbooks for the on-call rotation; cut median MTTR from 47 to 18 minutes.",
        "Set up Datadog dashboards and SLOs for the routing service; standardized monitoring patterns across the platform group.",
        "Refactored a monolithic dispatch service into three event-driven services; cut deploy lead time 4x.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Northwind Foods — Backend Engineer", level=2)
    doc.add_paragraph("2019 – 2023")
    for bullet in [
        "Built Postgres-backed inventory and pricing services in Python and FastAPI.",
        "Tuned Postgres query plans and indexing for the catalog service; improved key endpoint throughput 3x.",
        "Implemented webhook integrations with supplier APIs; handled idempotency and retry logic at the application layer.",
        "Migrated legacy Django reports onto a Looker dashboard; freed two engineers from one-off report requests.",
        "Wrote pytest integration suites with real Postgres in CI; cut prod regressions in the catalog service to near zero.",
        "Partnered with the data team to ship a daily ETL into Snowflake; replaced a fragile cron-on-EC2 job.",
        "Refactored CI/CD scripts from Bash into Python; made them testable.",
        "Reviewed PRs across three sister teams; informally mentored two engineers who later took senior roles.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Carver Robotics — Junior Backend Engineer", level=2)
    doc.add_paragraph("2017 – 2019")
    for bullet in [
        "Wrote Python services for the robot fleet's command-and-control layer.",
        "Built REST APIs in Flask for the operator console.",
        "Maintained ROS integration scripts that bridged the Python services and the robot firmware layer.",
        "Wrote CAD-pipeline tooling in Python to ingest part specs from the mechanical engineering team.",
        "Built test harnesses simulating robot sensor input for the backend.",
        "Contributed to monitoring dashboards for the fleet-management service.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "Python, Postgres, AWS (ECS, RDS, S3), Kafka, FastAPI, Flask, "
        "Datadog, Terraform, GitHub Actions, pytest, Docker."
    )

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. Computer Science, University of Washington, 2017")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  ✓ wrote {path.relative_to(REPO)}")

    # Also drop a copy into resumes/<DEMO_USER>/ so
    # corpus_import.run_import (which reads from that directory under
    # with_llm=true) finds it during the corpus seed.
    seed_dir = REPO / "resumes" / DEMO_USER
    seed_dir.mkdir(parents=True, exist_ok=True)
    seed_path = seed_dir / "priya_master.docx"
    shutil.copyfile(str(path), str(seed_path))
    print(f"  ✓ wrote {seed_path.relative_to(REPO)}")


def cap(page: Page, filename: str) -> None:
    out = SHOTS / filename
    out.parent.mkdir(parents=True, exist_ok=True)
    page.screenshot(path=str(out), full_page=False)
    print(f"  📸 {out.relative_to(REPO)}")


def wait_quiet(page: Page, ms: int = 600) -> None:
    """Small settle pause so animated transitions complete before capture."""
    page.wait_for_timeout(ms)


def ensure_demo_user(page: Page) -> None:
    """Create/select the demo user via the user-picker POM."""
    picker = UserPickerPage(page, APP_URL)
    if DEMO_USER in picker.options():
        print(f"  · {DEMO_USER!r} user already exists; selecting it")
        picker.select(DEMO_USER)
        return
    print(f"  · creating {DEMO_USER!r} user")
    picker.create(DEMO_USER, "Priya Sharma", "priya.sharma@example.com")


def ensure_corpus_imported(page: Page) -> None:
    """Import the synthetic .docx into the demo user's corpus.

    Calls onboarding.corpus_import.run_import directly (bypassing HTTP)
    to seed the demo user's DB corpus from resumes/<user>/.
    """
    corpus = CorpusPage(page, APP_URL).open()
    wait_quiet(page)

    # If the corpus is already populated (re-run with --keep-user),
    # skip the import entirely.
    existing = corpus.card_count()
    if existing > 0:
        print(f"  · corpus already has {existing} experiences; skipping import")
        cap(page, "walkthrough_setup_corpus-empty.png")
        return

    # Capture S03 — empty corpus state — BEFORE the import.
    cap(page, "walkthrough_setup_corpus-empty.png")

    print("  · running corpus_import.run_import(demo, with_llm=True)…")
    # ~$0.02 Haiku call for extraction; ~10-20s.
    # We call run_import directly (same path the Flask route wraps)
    # rather than going through HTTP, because the HTTP path was
    # failing with an opaque KeyError('config') that didn't reproduce
    # via curl or a REPL — bypassing it removes Playwright's request
    # encoding, the Flask route, and Flask debug-mode's auto-reload
    # window from the dependency chain.
    from onboarding.corpus_import import run_import

    report = run_import(DEMO_USER, with_llm=True)
    if report.errors:
        raise RuntimeError(f"corpus_import reported errors: {report.errors}")
    print(
        f"  · imported: {report.experiences_created} experiences "
        f"({report.experiences_merged} merged), "
        f"{report.bullets_created} bullets"
    )

    # Refresh the Corpus tab UI to render the freshly-imported rows.
    # The cache key _corpusLoadedForUser was set during the empty-state
    # render — clear it so refreshCorpus() actually re-fetches; wrap in
    # an async IIFE so page.evaluate awaits the refresh's promise.
    # Page reload so the freshly-populated DB state is consistent in
    # the page's JS context for the rest of the run. We don't need
    # the Corpus tab to actually render cards (Step 1+ reads from
    # backend APIs that hit the DB directly); the reload is purely
    # belt-and-braces so currentUser + corpus loaded state are clean.
    print("  · reloading page so corpus state is fresh in the UI")
    page.reload()
    UserPickerPage(page, APP_URL).select(DEMO_USER)
    wait_quiet(page, 1500)


def run_step1(page: Page) -> None:
    """Step 1 — paste JD, capture pre + post analyze."""
    job = WizardJobPage(page, APP_URL).open()
    wait_quiet(page)

    job.fill_jd(PRIYA_JD)
    wait_quiet(page)
    cap(page, "walkthrough_step1pre_jd-textarea.png")

    print("  · clicking Analyze; this is the ~30-60s Sonnet 5 call…")
    job.analyze()
    wait_quiet(page, 1200)
    page.evaluate("""
        const topbarH = document.querySelector('#cbTopbar').getBoundingClientRect().height;
        document.querySelector('#panelAnalysis .panel-header').scrollIntoView({block: 'start'});
        window.scrollBy(0, -topbarH);
    """)
    wait_quiet(page, 300)
    cap(page, "walkthrough_step1post_analysis-filled.png")
    # Same state, hero shot for README.
    cap(page, "readme_hero_wizard-step1-filled.png")


def run_step2(page: Page) -> None:
    """Step 2 — get clarify questions, type a partial answer.

    Forward navigation uses the in-flow "Continue to Clarify →"
    button, not the wizard rail. The rail's step 2 button stays
    disabled (class=upcoming) until _wizardRender is invoked, which
    only happens via wizardGoTo() — clicking the rail directly is
    a no-op when the step isn't yet "reached." The in-flow button
    calls wizardGoTo(2), which then enables the rail.
    """
    WizardJobPage(page, APP_URL).continue_to_clarify()
    wait_quiet(page)

    clarify = WizardClarifyPage(page, APP_URL)
    print("  · waiting on clarification questions (~30s Sonnet call)…")
    # "Continue to Clarify →" already auto-fetches (app.js finding #6) —
    # #btnClarify is disabled for the duration, so wait rather than click.
    clarify.wait_for_questions()
    wait_quiet(page, 1000)

    # Type a partial answer in the first question so the capture
    # shows realistic mid-typing state.
    clarify.answer_first(
        "Lead implementer. Designed the topic + partition scheme — 12 topics, "
        "60 partitions on the busiest one,",
    )
    wait_quiet(page)
    page.evaluate("""
        const topbarH = document.querySelector('#cbTopbar').getBoundingClientRect().height;
        document.querySelector('#panelClarify .panel-header').scrollIntoView({block: 'start'});
        window.scrollBy(0, -topbarH);
    """)
    wait_quiet(page, 300)
    cap(page, "walkthrough_step2_clarify-questions.png")


def run_step3(page: Page) -> None:
    """Step 3 — Compose. Wait for cards, capture; no manual pinning."""
    # Submit minimal clarifications by filling all blanks, then continue.
    clarify = WizardClarifyPage(page, APP_URL)
    clarify.fill_blank_answers("Yes — see prior answer for related context.")
    clarify.submit_to_compose()
    wait_quiet(page)

    print("  · waiting for compose recommendations (Haiku, multiple calls)…")
    WizardComposePage(page, APP_URL).wait_cards()
    wait_quiet(page, 1500)
    cap(page, "walkthrough_step3_compose-experience-card.png")


def run_step4(page: Page) -> None:
    """Step 4 — Template. Modern is the default; let live preview render."""
    WizardComposePage(page, APP_URL).continue_to_template()
    wait_quiet(page)
    # Pick the second card by index as a stable proxy for "Modern"
    # (Classic is usually first), then let the live preview render.
    template = WizardTemplatePage(page, APP_URL)
    template.pick_template(1)
    template.wait_live_preview()
    wait_quiet(page, 2000)
    cap(page, "walkthrough_step4_template-modern-preview.png")


def run_step5_and_6(page: Page) -> None:
    """Step 5 — Generate. Step 6 — Download. Capture Refine state."""
    WizardTemplatePage(page, APP_URL).continue_to_generate()
    wait_quiet(page)

    print("  · generating résumé (~30-60s Sonnet 5 call)…")
    gen = WizardGeneratePage(page, APP_URL)
    gen.generate()
    wait_quiet(page, 2000)

    # Type a sample refine note so the capture shows what refinement looks
    # like in practice (no-op if the input isn't present yet).
    if gen.refine(SAMPLE_REFINE_NOTE):
        wait_quiet(page)
    else:
        print("  · refinement input not visible; capturing without it")
    cap(page, "walkthrough_step6_download-with-refine.png")


def run_cover_letter(page: Page) -> None:
    """Optional — generate cover letter, capture."""
    print("  · generating cover letter (~30s Sonnet call)…")
    if not WizardOutputPage(page, APP_URL).generate_cover_letter():
        print("  · cover-letter button not found; skipping S10")
        return
    wait_quiet(page, 1500)
    cap(page, "walkthrough_coverletter_first-generation.png")


def dismiss_welcome_modal_if_present(page: Page) -> None:
    """Dismiss the once-ever first-view welcome help modal (static/app.js's
    `_HELP_REGISTRY` auto-open), which opens unconditionally on any fresh,
    empty-localStorage browser context — exactly what this script's
    Playwright context always is. Same click-away dismissal the UX suite
    already uses (tests/ux/regression/test_20260614_help_pattern.py).
    No-op if the modal isn't present (e.g. a --keep-user re-run).
    """
    modal = page.locator(S.Help.MODAL)
    if modal.is_visible():
        page.click(S.Help.BACKDROP, position={"x": 5, "y": 5})
        page.wait_for_selector(S.Help.MODAL, state="hidden", timeout=SHORT_TIMEOUT_MS)


def capture_user_picker(page: Page) -> None:
    """S02 — capture the user-picker section with the New User form open.

    Native <select> dropdowns are OS-rendered and not Playwright-screenshot-
    able while open, so we use the New User form (DOM-rendered) as a more
    informative substitute. Shows both the picker and the create affordance.
    """
    page.goto(APP_URL)
    page.wait_for_selector(S.UserPicker.PANEL, state="visible", timeout=SHORT_TIMEOUT_MS)
    dismiss_welcome_modal_if_present(page)
    page.click(S.UserPicker.NEW_USER_LINK)
    page.wait_for_selector(S.UserPicker.NEW_USER_FORM, state="visible")
    page.fill(S.UserPicker.NEW_USERNAME, "your-username")
    wait_quiet(page)
    cap(page, "install_setup_user-picker.png")
    # Restore a clean state for the rest of the run.
    page.reload()


def cleanup(keep_user: bool) -> None:
    if keep_user:
        print(f"  · keeping {DEMO_USER!r} user + corpus per --keep-user flag")
        return
    print("  · cleaning up demo user artefacts…")
    for sub in [
        REPO / "configs" / f"{DEMO_USER}.config",
        REPO / "resumes" / DEMO_USER,
        REPO / "output" / DEMO_USER,
        REPO / "personas" / "owned" / DEMO_USER,
        TMP,
    ]:
        if sub.exists():
            if sub.is_dir():
                shutil.rmtree(sub, ignore_errors=True)
            else:
                sub.unlink(missing_ok=True)
            print(f"    · removed {sub.relative_to(REPO)}")
    # Note: the demo user's rows remain in db/resume.sqlite. That's fine
    # for re-runs; the user can drop the row manually if they want a
    # truly clean DB.


def main() -> int:
    # Windows cp1252 consoles can't encode this script's unicode progress output
    # (the ✓/📸 markers, em-dashes, ellipses) or the em-dashes in `--help`/`__doc__`
    # — force UTF-8 on our streams before argparse or any print runs (the EV-3
    # crash class, window-8.5-findings).
    for _stream in (sys.stdout, sys.stderr):
        # suppress on a non-reconfigurable stream (e.g. piped)
        with contextlib.suppress(AttributeError, ValueError):
            _stream.reconfigure(encoding="utf-8")  # type: ignore[union-attr]

    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--headless", action="store_true", help="run Chromium headless")
    ap.add_argument(
        "--keep-user",
        action="store_true",
        help="don't delete the demo user + on-disk artefacts after capture",
    )
    ap.add_argument(
        "--smoke",
        action="store_true",
        help=(
            "run setup + Step 1 only, then stop (periodic drift smoke check; "
            "skips the Step 2-6 + cover-letter LLM calls)"
        ),
    )
    args = ap.parse_args()

    SHOTS.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    docx_path = TMP / "priya_master.docx"

    print("=" * 70)
    print("sartor. screenshot capture")
    print("=" * 70)
    print(f"App URL:     {APP_URL}")
    print(f"Demo user:   {DEMO_USER}")
    print(f"Viewport:    {VIEWPORT['width']}x{VIEWPORT['height']}")
    print(f"Output dir:  {SHOTS.relative_to(REPO)}")
    print()

    print("1) Generating synthetic Priya master .docx")
    write_priya_docx(docx_path)
    print()

    start = time.time()
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=args.headless)
        ctx = browser.new_context(
            viewport=VIEWPORT,  # type: ignore[arg-type]
            color_scheme="light",
        )
        page = ctx.new_page()
        page.set_default_timeout(SHORT_TIMEOUT_MS)
        # Suppress every auto-firing help block (welcome + KW3 tour stops +
        # dashboard explainers) before any navigation — this script always
        # gets a fresh, empty-localStorage context and always creates a
        # brand-new demo user, so every auto-open condition would otherwise
        # fire at some point during the walkthrough. See
        # docs/dev/diagnosis/capture-screenshots-welcome-modal.md.
        page.add_init_script(S.Help.suppress_tour_init_script())

        print("2) Capturing S02 — user picker section")
        capture_user_picker(page)
        print()

        print("3) Creating demo user")
        page.goto(APP_URL)
        ensure_demo_user(page)
        print()

        print("4) Capturing S03 — empty corpus + importing Priya .docx")
        ensure_corpus_imported(page)
        print()

        print("5) Step 1 — Job + Analyze (S01 hero + S04 pre + S05 post)")
        run_step1(page)
        print()

        if args.smoke:
            print("--smoke: stopping after Step 1 (periodic drift smoke check)")
            print()
        else:
            print("6) Step 2 — Clarify (S06)")
            run_step2(page)
            print()

            print("7) Step 3 — Compose (S07)")
            run_step3(page)
            print()

            print("8) Step 4 — Template (S08)")
            run_step4(page)
            print()

            print("9) Steps 5 + 6 — Generate + Download (S09)")
            run_step5_and_6(page)
            print()

            print("10) Cover letter (S10)")
            run_cover_letter(page)
            print()

        browser.close()

    elapsed = time.time() - start
    print(f"✓ capture complete in {elapsed:.1f}s")
    print()

    cleanup(keep_user=args.keep_user)

    print()
    print("Next steps:")
    print(f"  · review {SHOTS.relative_to(REPO)}/")
    print("  · if any capture looks off, re-run with --keep-user and iterate")
    print("  · once the 10 PNGs look right, I (Claude) can do the markdown")
    print("    insertion pass per docs/ux/screenshot_capture.md")
    return 0


if __name__ == "__main__":
    sys.exit(main())
