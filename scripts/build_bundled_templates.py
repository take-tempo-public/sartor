"""Generate the 4 bundled ATS-friendly persona templates.

Phase C.1 / v1.0.0 curation: produces
`personas/bundled/{classic,modern,spacious,tech}.docx`
programmatically via python-docx. The templates are STYLE CARRIERS — they
contain placeholder paragraphs in the role order that `generator.py:_capture_template_styles`
captures (name, subtitle, contact, section_heading, job_title, job_subtitle,
body, bullet). Each template differs in typography + spacing, not structure.

ATS rules (see docs/template_authoring.md):
- Single column only; no tables, text boxes, or images
- Standard fonts (Arial, Calibri, Helvetica); 11pt body, 12-14pt headings
- Standard section headings (Experience, Education, Skills, ...)
- Bullet glyphs: `•` or `-` (we use `-` via Word's List Bullet style)
- Right-aligned dates via tab stop on job_title lines

Re-run anytime to regenerate. Idempotent — overwrites existing files.

Usage:
    python -m scripts.build_bundled_templates
    python -m scripts.build_bundled_templates --out personas/bundled
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT = REPO_ROOT / "personas" / "bundled"

# Standard right-margin tab position for the date column on job_title lines.
# 6.5in matches a US-letter page with 1in margins; templates with tighter
# margins override via _build_template_for_preset.
DATE_TAB_INCHES = 6.5

# License notice we burn into every bundled .docx's core properties so the
# user can inspect provenance via Word's File → Info pane.
LICENSE_NOTICE = (
    "MIT-licensed bundled persona template. See docs/bundled_templates_LICENSE.md "
    "for license text and credits."
)


# ---------------------------------------------------------------------------
# Style preset model — one per template
# ---------------------------------------------------------------------------


@dataclass
class TypographyPreset:
    """All the typography knobs each template varies. Structure is identical
    across templates; only these values differ. Keeps the build script
    declarative: add a new template = add a new preset."""

    filename: str
    display_name: str
    description: str

    # Fonts + sizes
    font_family: str = "Calibri"
    name_pt: int = 16
    subtitle_pt: int = 11
    contact_pt: int = 10
    section_heading_pt: int = 12
    job_title_pt: int = 11
    job_subtitle_pt: int = 10
    body_pt: int = 11
    bullet_pt: int = 11

    # Spacing (in points; 1 line = 12pt × line_spacing)
    line_spacing: float = 1.15
    body_space_after_pt: int = 6
    section_space_before_pt: int = 12
    section_space_after_pt: int = 6
    bullet_space_after_pt: int = 3

    # Page margins
    margin_inches: float = 1.0

    # Visual flourishes (all single-column, ATS-safe; only typography)
    name_alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER
    section_heading_uppercase: bool = False
    section_heading_underline: bool = False
    section_heading_small_caps: bool = False

    # Sample content the template ships with (placeholders for the user to overwrite)
    sample_name: str = "Casey Rivera"
    sample_subtitle: str = "Senior Product Manager"
    sample_contact: str = "casey.rivera@example.com | 555-0142 | linkedin.com/in/casey-rivera"

    # The persona's tag suggestions for the seed migration (kept here so the
    # build script and the migration agree on the canonical mapping)
    suggested_role_tags: list[str] = field(default_factory=list)


PRESETS: list[TypographyPreset] = [
    TypographyPreset(
        filename="classic.docx",
        display_name="Classic Single-Column",
        description=(
            "Maximally ATS-safe baseline. Arial 11pt, conservative spacing, "
            "uppercase section headings. The default fallback when no "
            "candidate-chosen persona matches a JD's role tag."
        ),
        font_family="Arial",
        name_pt=16,
        section_heading_pt=12,
        section_heading_uppercase=True,
        line_spacing=1.15,
        body_space_after_pt=6,
        section_space_before_pt=12,
        suggested_role_tags=["generalist"],
    ),
    TypographyPreset(
        filename="modern.docx",
        display_name="Modern Single-Column",
        description=(
            "Calibri 11pt with mild typographic refinement: small-caps section "
            "headings, tighter line spacing. Good middle-ground for most roles."
        ),
        font_family="Calibri",
        name_pt=15,
        section_heading_pt=11,
        section_heading_small_caps=True,
        line_spacing=1.1,
        body_space_after_pt=4,
        section_space_before_pt=10,
        suggested_role_tags=["pm", "design-mgmt", "generalist"],
    ),
    TypographyPreset(
        filename="spacious.docx",
        display_name="Spacious (Career Changer / Junior)",
        description=(
            "Arial 11pt with generous spacing. Built for early-career or "
            "career-changing candidates with less to fit, prioritizes readability."
        ),
        font_family="Arial",
        name_pt=17,
        section_heading_pt=13,
        section_heading_uppercase=True,
        line_spacing=1.5,
        body_space_after_pt=10,
        section_space_before_pt=16,
        section_space_after_pt=10,
        bullet_space_after_pt=4,
        margin_inches=1.0,
        suggested_role_tags=["ic-design", "generalist"],
    ),
    TypographyPreset(
        filename="tech.docx",
        display_name="Tech (ATS-optimized)",
        description=(
            "Georgia 11pt with centered name and underlined section headings. "
            "Single column, plain bullets, no inline glyphs — designed for "
            "engineering / data / AI roles where the parser must catch every "
            "tech keyword verbatim. Inspired by the community "
            "jsonresume-theme-dev-ats."
        ),
        font_family="Georgia",
        name_pt=15,
        section_heading_pt=12,
        section_heading_underline=True,
        line_spacing=1.2,
        body_space_after_pt=5,
        section_space_before_pt=10,
        suggested_role_tags=["ai", "engineering", "data-science", "physical-compute"],
    ),
]


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------


def _apply_font(
    run,
    family: str,
    size_pt: int,
    *,
    bold: bool = False,
    italic: bool = False,
    small_caps: bool = False,
) -> None:
    """Set typography on a single run."""
    run.font.name = family
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if small_caps:
        run.font.small_caps = True


def _apply_paragraph_spacing(
    paragraph,
    *,
    line_spacing: float,
    space_before_pt: int = 0,
    space_after_pt: int = 0,
) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


def build_template_to_path(preset: TypographyPreset, out_path: Path) -> None:
    """Generate one .docx template at the given path."""
    doc = Document()

    # Set page margins per preset
    for section in doc.sections:
        section.left_margin = Inches(preset.margin_inches)
        section.right_margin = Inches(preset.margin_inches)
        section.top_margin = Inches(preset.margin_inches)
        section.bottom_margin = Inches(preset.margin_inches)

    # Core properties — the MIT notice the plan calls for
    doc.core_properties.title = preset.display_name
    doc.core_properties.subject = "Resume template (MIT-licensed)"
    doc.core_properties.comments = LICENSE_NOTICE
    doc.core_properties.category = preset.display_name

    # 1) Name (role: "name")
    p_name = doc.add_paragraph()
    p_name.alignment = preset.name_alignment
    _apply_paragraph_spacing(p_name, line_spacing=1.0, space_after_pt=0)
    run = p_name.add_run(preset.sample_name)
    _apply_font(run, preset.font_family, preset.name_pt, bold=True)

    # 2) Subtitle (role: "subtitle")
    p_sub = doc.add_paragraph()
    p_sub.alignment = preset.name_alignment
    _apply_paragraph_spacing(p_sub, line_spacing=1.0, space_after_pt=0)
    run = p_sub.add_run(preset.sample_subtitle)
    _apply_font(run, preset.font_family, preset.subtitle_pt)

    # 3) Contact (role: "contact")
    p_contact = doc.add_paragraph()
    p_contact.alignment = preset.name_alignment
    _apply_paragraph_spacing(
        p_contact, line_spacing=1.0, space_after_pt=preset.section_space_before_pt
    )
    run = p_contact.add_run(preset.sample_contact)
    _apply_font(run, preset.font_family, preset.contact_pt)

    # 4) First section heading (role: "section_heading")
    _add_section_heading(doc, preset, "Experience")

    # 5) Job title line with right-aligned date tab (role: "job_title")
    p_job = doc.add_paragraph()
    _apply_paragraph_spacing(
        p_job,
        line_spacing=preset.line_spacing,
        space_before_pt=preset.section_space_after_pt,
        space_after_pt=0,
    )
    # Right tab stop for the date column
    tab_pos = (preset.margin_inches and Inches(8.5 - 2 * preset.margin_inches)) or Inches(
        DATE_TAB_INCHES
    )
    p_job.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
    run = p_job.add_run("Sample Company, Senior Role")
    _apply_font(run, preset.font_family, preset.job_title_pt, bold=True)
    p_job.add_run("\tJanuary 2022 – Present").font.size = Pt(preset.job_title_pt)
    for r in p_job.runs:
        r.font.name = preset.font_family

    # 6) Job subtitle (role: "job_subtitle") — optional context line
    p_sub2 = doc.add_paragraph()
    _apply_paragraph_spacing(
        p_sub2,
        line_spacing=preset.line_spacing,
        space_after_pt=preset.body_space_after_pt,
    )
    run = p_sub2.add_run("Optional one-line context (team, scope, technology)")
    _apply_font(run, preset.font_family, preset.job_subtitle_pt, italic=True)

    # 7) Body paragraph (role: "body")
    p_body = doc.add_paragraph()
    _apply_paragraph_spacing(
        p_body,
        line_spacing=preset.line_spacing,
        space_after_pt=preset.body_space_after_pt,
    )
    run = p_body.add_run(
        "Plain body paragraphs use this style — typical resume practice favors bullets, but some sections (Summary) read better as prose."
    )
    _apply_font(run, preset.font_family, preset.body_pt)

    # 8) Bullets (role: "bullet")
    for text in (
        "Sample bullet — strong action verb up front, measurable outcome where possible.",
        "Second sample bullet — preserved metrics, no fabricated specifics.",
        "Third sample bullet — JD-keyword integration without invention.",
    ):
        p_bul = doc.add_paragraph(style="List Bullet")
        _apply_paragraph_spacing(
            p_bul,
            line_spacing=preset.line_spacing,
            space_after_pt=preset.bullet_space_after_pt,
        )
        run = p_bul.add_run(text)
        _apply_font(run, preset.font_family, preset.bullet_pt)

    # 9) Additional sections so the template demonstrates the full role set
    _add_section_heading(doc, preset, "Education")
    p_edu = doc.add_paragraph()
    _apply_paragraph_spacing(
        p_edu,
        line_spacing=preset.line_spacing,
        space_after_pt=preset.body_space_after_pt,
    )
    run = p_edu.add_run("BS Cognitive Science, ExampleU (2010–2014)")
    _apply_font(run, preset.font_family, preset.body_pt)

    _add_section_heading(doc, preset, "Skills")
    p_skills = doc.add_paragraph()
    _apply_paragraph_spacing(
        p_skills,
        line_spacing=preset.line_spacing,
        space_after_pt=preset.body_space_after_pt,
    )
    run = p_skills.add_run("Sample skill 1 · Sample skill 2 · Sample skill 3 · Sample skill 4")
    _apply_font(run, preset.font_family, preset.body_pt)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _add_section_heading(doc, preset: TypographyPreset, text: str) -> None:
    p = doc.add_paragraph()
    _apply_paragraph_spacing(
        p,
        line_spacing=1.0,
        space_before_pt=preset.section_space_before_pt,
        space_after_pt=preset.section_space_after_pt,
    )
    display_text = text.upper() if preset.section_heading_uppercase else text
    run = p.add_run(display_text)
    _apply_font(
        run,
        preset.font_family,
        preset.section_heading_pt,
        bold=True,
        small_caps=preset.section_heading_small_caps,
    )
    if preset.section_heading_underline:
        run.underline = True


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    parser.add_argument(
        "--out",
        default=str(DEFAULT_OUT),
        help="Output directory (defaults to personas/bundled at repo root)",
    )
    args = parser.parse_args(argv)
    out_dir = Path(args.out)

    print(f"Building {len(PRESETS)} bundled templates -> {out_dir}")
    for preset in PRESETS:
        out_path = out_dir / preset.filename
        build_template_to_path(preset, out_path)
        print(f"  [ok] {preset.filename:24s} ({preset.display_name})")
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
