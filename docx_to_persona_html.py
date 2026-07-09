"""Generate an HTML+CSS preview companion for an uploaded persona .docx.

Deterministic — no LLM (charter **C-6**). The live résumé preview
(`blueprints/templates.py:preview_application_html`) renders a JSON Resume
document through a Jinja2 `.html` + `.css` companion resolved as the sibling of
the chosen `.docx` (`pdf_render.html_template_path_for`). The 4 bundled personas
ship hand-authored companions; **uploaded** user `.docx` templates do not, so the
preview silently fell back to `classic.html` — every uploaded template previewed
as Classic 1-column even though the `.docx` download was faithful (walkthrough
B2 / B3 / Step-6 #4).

This module closes that gap: at upload time (and lazily on first preview for
personas uploaded before this shipped) it reads the uploaded `.docx` with
python-docx, extracts the typography that matters for a single-column preview
(page margins, base font family/size, name / heading / job-title sizes, heading
treatment — uppercase / small-caps / underline / color, line spacing), and emits
a companion `.html` (a verbatim copy of the canonical `classic.html` skeleton
with the CSS `href` swapped — so the Jinja2 contract consumed by
`pdf_render.render_html_string` is byte-identical) plus a `.css` derived from the
extracted knobs (Classic's ATS-safe single-column structure, re-typed).

Fidelity ceiling: python-docx cannot faithfully represent multi-column sections,
tables, text boxes, or floating images. We ALWAYS render single-column (matching
what `generator._write_docx` produces on download, so preview == download stay
mutually consistent), and record an honest `layout_fidelity` marker
(`full` vs `typography_only`) in a `<stem>.persona.json` sidecar so callers/logs
know when the preview is a typographic approximation of a fancier source.

The role-classification heuristics mirror `generator._capture_template_styles`
(centered top-3 = name/subtitle/contact; bold + right-tab = job_title; bold =
section_heading; the non-bold after a job_title = job_subtitle; else body) so the
preview reads the .docx the same way the download does.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Length
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)

# The canonical HTML skeleton + CSS reference model. This module lives at the
# repo root, so its parent IS the repo root.
_REPO_ROOT = Path(__file__).resolve().parent
_SKELETON_HTML = _REPO_ROOT / "personas" / "bundled" / "classic.html"

# Serif families we recognize; everything else maps to a sans stack. Used only to
# pick a sensible CSS fallback chain + generic family for the extracted font.
_SERIF_FAMILIES = {
    "georgia",
    "times new roman",
    "times",
    "cambria",
    "garamond",
    "book antiqua",
    "palatino",
    "liberation serif",
}

# Neutral defaults when the .docx leaves an attribute unset (inherited / default).
_DEFAULT_FG = "#111"
_DEFAULT_ACCENT = "#b87333"  # Classic's warm bronze rule; prints well in B&W.


def _css_font_stack(family: str | None) -> str:
    """Map a single docx font name to a CSS font stack with a generic fallback."""
    fam = (family or "").strip()
    if not fam:
        return '"Helvetica Neue", Helvetica, Arial, "Liberation Sans", sans-serif'
    if fam.lower() in _SERIF_FAMILIES:
        return f'"{fam}", Georgia, "Times New Roman", Cambria, "Liberation Serif", serif'
    return f'"{fam}", "Helvetica Neue", Helvetica, Arial, "Liberation Sans", sans-serif'


def _run_color_hex(run: Run) -> str | None:
    """Return an explicit RGB run color as `#rrggbb`, or None if inherited/theme."""
    try:
        color = run.font.color
        if color is None or color.type is None:
            return None
        rgb = color.rgb
    except (AttributeError, ValueError):
        return None
    if rgb is None:
        return None
    return f"#{rgb!s}".lower()


def _has_right_tab(paragraph: Paragraph) -> bool:
    return any(
        t.alignment == WD_TAB_ALIGNMENT.RIGHT for t in (paragraph.paragraph_format.tab_stops or [])
    )


def detect_layout_fidelity(doc: DocxDocument) -> str:
    """Classify whether the .docx is a single-column text layout we can render.

    Returns `"full"` when the source is single-column running text (our preview
    matches faithfully) or `"typography_only"` when it contains tables, multiple
    columns, text boxes, or drawings/images — layout python-docx can't represent,
    so the preview carries the source's typography on a clean single column.
    """
    if doc.tables:
        return "typography_only"
    for section in doc.sections:
        cols = section._sectPr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"))
            if num is not None and num.isdigit() and int(num) > 1:
                return "typography_only"
    body_xml = doc.element.body.xml
    if "txbxContent" in body_xml or "w:drawing" in body_xml or "pic:pic" in body_xml:
        return "typography_only"
    return "full"


def extract_persona_style(docx_path: str | Path) -> dict[str, Any]:
    """Read an uploaded persona .docx and return the typography knob dict.

    Deterministic. Mirrors `generator._capture_template_styles`'s role heuristics
    so the preview interprets the .docx the same way the .docx download does.
    Missing attributes fall back to neutral defaults so a partial/odd .docx still
    yields a clean companion.

    Also captures per-role vertical rhythm (mirroring `generator._capture_proto`'s
    `space_before_pt`/`space_after_pt`, not just font sizes): the header block's
    closing gap (`header_space_after_pt`, from the LAST centered header line —
    "contact" when a 3-line header exists, "subtitle" when it's a 2-line header),
    the section heading's `heading_space_before_pt`/`heading_space_after_pt`, and
    `job_title_space_before_pt` (the gap before each job entry). Each stays `None`
    when the source `.docx` never sets it explicitly (inherits from the style),
    so `_build_css` keeps its existing hardcoded default — no behavior change for
    documents without direct paragraph-level spacing overrides.

    `job_title_has_right_tab` records whether the job-title line actually carries
    a right-aligned tab stop for the date column. Real uploaded résumés routinely
    don't (the date is a bare tab character with no defined stop — Word falls
    through to its default tab ladder, NOT a flush-right column); `_build_css`
    reads this knob so the preview doesn't idealize a right-alignment the actual
    `.docx` doesn't render (owner decision: no right-alignment work in the docx
    writer — the preview reconciles toward the document, not the other way
    around).
    """
    doc = Document(str(docx_path))

    section = doc.sections[0] if doc.sections else None

    def _margin_in(value: Length | None, default: float) -> float:
        try:
            return round(value.inches, 3) if value is not None else default
        except (AttributeError, TypeError):
            return default

    knobs: dict[str, Any] = {
        "font_family": None,
        "base_font_pt": 11.0,
        "margin_top_in": _margin_in(getattr(section, "top_margin", None), 1.0),
        "margin_right_in": _margin_in(getattr(section, "right_margin", None), 1.0),
        "margin_bottom_in": _margin_in(getattr(section, "bottom_margin", None), 1.0),
        "margin_left_in": _margin_in(getattr(section, "left_margin", None), 1.0),
        "name_pt": 22.0,
        "name_align": "center",
        "subtitle_pt": 11.0,
        "contact_pt": 9.5,
        "heading_pt": 12.0,
        "heading_uppercase": False,
        "heading_small_caps": False,
        "heading_underline": False,
        "heading_color": None,
        "line_spacing": 1.3,
        "job_title_pt": 11.0,
        "job_title_has_right_tab": False,
        "header_space_after_pt": None,
        "heading_space_before_pt": None,
        "heading_space_after_pt": None,
        "job_title_space_before_pt": None,
        "layout_fidelity": detect_layout_fidelity(doc),
    }

    centered_seen = 0
    last_was_job_title = False
    seen_heading = False
    seen_body = False

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        run0 = p.runs[0] if p.runs else None
        size_pt = run0.font.size.pt if (run0 and run0.font.size) else None
        family = run0.font.name if (run0 and run0.font.name) else None
        if family and not knobs["font_family"]:
            knobs["font_family"] = family
        is_centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        any_bold = any(r.bold for r in p.runs)
        has_right_tab = _has_right_tab(p)
        line_spacing = p.paragraph_format.line_spacing

        # Header zone: centered paragraphs at the top, in order.
        if is_centered and centered_seen < 3 and not seen_heading:
            role = ("name", "subtitle", "contact")[centered_seen]
            if size_pt:
                knobs[f"{role}_pt"] = size_pt
            if role == "name":
                knobs["name_align"] = "center"
            # Overwritten on every header line seen so it ends on the LAST one —
            # "contact" for a 3-line header, "subtitle" for the common 2-line
            # case (name + contact, no separate tagline).
            header_sa = p.paragraph_format.space_after
            if header_sa is not None:
                knobs["header_space_after_pt"] = header_sa.pt
            centered_seen += 1
            continue

        # Bold + right tab stop = job title with right-aligned date.
        if any_bold and has_right_tab:
            if size_pt:
                knobs["job_title_pt"] = size_pt
            knobs["job_title_has_right_tab"] = True
            job_sb = p.paragraph_format.space_before
            if job_sb is not None:
                knobs["job_title_space_before_pt"] = job_sb.pt
            last_was_job_title = True
            continue

        # Bold without tab stops = section heading — capture its treatment once.
        if any_bold:
            if not seen_heading:
                seen_heading = True
                if size_pt:
                    knobs["heading_pt"] = size_pt
                text = p.text.strip()
                knobs["heading_uppercase"] = (
                    bool(text) and text == text.upper() and any(c.isalpha() for c in text)
                )
                if run0 is not None:
                    knobs["heading_small_caps"] = bool(run0.font.small_caps)
                    knobs["heading_underline"] = bool(run0.font.underline or run0.underline)
                    knobs["heading_color"] = _run_color_hex(run0)
                heading_sb = p.paragraph_format.space_before
                if heading_sb is not None:
                    knobs["heading_space_before_pt"] = heading_sb.pt
                heading_sa = p.paragraph_format.space_after
                if heading_sa is not None:
                    knobs["heading_space_after_pt"] = heading_sa.pt
            last_was_job_title = False
            continue

        # Non-bold immediately after a job title = job subtitle; skip it so the
        # base font size below reflects the true body paragraph, not the subtitle.
        if last_was_job_title:
            last_was_job_title = False
            continue

        # First real body paragraph → base font size + line spacing.
        if not seen_body:
            seen_body = True
            if size_pt:
                knobs["base_font_pt"] = size_pt
            if isinstance(line_spacing, (int, float)):
                knobs["line_spacing"] = float(line_spacing)
        last_was_job_title = False

    return knobs


def _fmt_pt(value: float) -> str:
    """Render a point size without a trailing `.0` (11.0 → '11', 10.5 → '10.5')."""
    return f"{value:g}"


def _spacing_css(value_pt: float | None, default_px: float) -> str:
    """Render a captured `space_before`/`space_after` knob as CSS.

    Falls back to the historical hardcoded `<default_px>px` literal when the
    source `.docx` never set the property explicitly (`None` — inherited from
    the paragraph style, not zero), so a document with no direct spacing
    overrides renders byte-identical CSS to before this knob existed.
    """
    if value_pt is None:
        return f"{_fmt_pt(default_px)}px"
    return f"{_fmt_pt(value_pt)}pt"


def _build_css(knobs: dict[str, Any]) -> str:
    """Emit a single-column ATS-safe stylesheet from the extracted knobs.

    Structure mirrors `personas/bundled/classic.css`; only the typography values
    are swapped for the uploaded template's own. Vertical rhythm (header/heading/
    job-entry gaps) uses the captured per-role spacing when the source `.docx` set
    it explicitly, else the same literals `classic.css` has always used.
    """
    fg = knobs.get("heading_color") or _DEFAULT_FG
    accent = knobs.get("heading_color") or _DEFAULT_ACCENT
    font_stack = _css_font_stack(knobs.get("font_family"))
    base_pt = _fmt_pt(knobs.get("base_font_pt", 11.0))
    line_spacing = knobs.get("line_spacing", 1.3)
    heading_transform = "uppercase" if knobs.get("heading_uppercase") else "none"
    heading_variant = "small-caps" if knobs.get("heading_small_caps") else "normal"
    heading_decoration = "underline" if knobs.get("heading_underline") else "none"

    header_space_after = _spacing_css(knobs.get("header_space_after_pt"), 20)
    heading_space_before = _spacing_css(knobs.get("heading_space_before_pt"), 20)
    heading_space_after = _spacing_css(knobs.get("heading_space_after_pt"), 8)
    job_space_after = _spacing_css(knobs.get("job_title_space_before_pt"), 14)

    # Date-column honesty (walkthrough residuals, item 1): only force the date
    # flush against the right margin when the source .docx actually right-aligns
    # it via a tab stop. Real uploaded résumés commonly leave the date as a bare
    # tab with no stop defined — Word renders that at its own default tab ladder,
    # NOT flush-right — so idealizing it here would show a layout the downloaded
    # .docx doesn't produce. Bundled personas all define the tab stop, so their
    # preview is unaffected (still `space-between`).
    has_date_tab = bool(knobs.get("job_title_has_right_tab"))
    job_header_justify = "space-between" if has_date_tab else "flex-start"
    job_header_gap = "8px" if has_date_tab else "18pt"

    m = (
        f"{_fmt_pt(knobs.get('margin_top_in', 1.0))}in "
        f"{_fmt_pt(knobs.get('margin_right_in', 1.0))}in "
        f"{_fmt_pt(knobs.get('margin_bottom_in', 1.0))}in "
        f"{_fmt_pt(knobs.get('margin_left_in', 1.0))}in"
    )

    return f"""\
/* Auto-generated preview companion for an uploaded persona .docx.
 * Deterministic output of docx_to_persona_html.generate_companion - do not
 * hand-edit; re-uploading the .docx regenerates it. Single-column + ATS-safe
 * (mirrors classic.css) re-typed with the uploaded template's typography.
 * layout_fidelity: {knobs.get("layout_fidelity", "full")}
 */

@page {{ size: letter; margin: {m}; }}

:root {{
  --fg-0: {fg};
  --fg-1: #333;
  --fg-2: #666;
  --rule: #ccc;
  --accent: {accent};
}}

* {{ margin: 0; padding: 0; box-sizing: border-box; }}

html, body {{
  font-family: {font_stack};
  font-size: {base_pt}pt;
  line-height: {line_spacing};
  color: var(--fg-0);
  background: #fff;
}}

.resume-header {{
  text-align: {knobs.get("name_align", "center")};
  margin-bottom: {header_space_after};
  padding-bottom: 14px;
  border-bottom: 1px solid var(--rule);
}}
.resume-header .name {{ font-size: {_fmt_pt(knobs.get("name_pt", 22.0))}pt; font-weight: 600; letter-spacing: 0.5pt; margin-bottom: 2pt; }}
.resume-header .label {{ font-size: {_fmt_pt(knobs.get("subtitle_pt", 11.0))}pt; color: var(--fg-1); margin-bottom: 4pt; font-weight: 400; }}
.resume-header .contact {{ font-size: {_fmt_pt(knobs.get("contact_pt", 9.5))}pt; color: var(--fg-2); letter-spacing: 0.1pt; }}

h2 {{
  font-size: {_fmt_pt(knobs.get("heading_pt", 12.0))}pt;
  font-weight: 600;
  color: var(--fg-0);
  text-transform: {heading_transform};
  font-variant: {heading_variant};
  text-decoration: {heading_decoration};
  letter-spacing: 1.2pt;
  margin-top: {heading_space_before};
  margin-bottom: {heading_space_after};
  padding-bottom: 2pt;
  border-bottom: 1.5pt solid var(--accent);
  page-break-after: avoid;
}}
section:first-of-type h2 {{ margin-top: 8px; }}

.summary p {{ line-height: 1.5; }}

.job, .degree, .project {{ margin-bottom: {job_space_after}; page-break-inside: avoid; }}
.job-header, .degree-header {{ display: flex; justify-content: {job_header_justify}; align-items: baseline; gap: {job_header_gap}; margin-bottom: 2pt; }}
.job-title, .degree-header h3 {{ font-size: {_fmt_pt(knobs.get("job_title_pt", 11.0))}pt; font-weight: 600; color: var(--fg-0); }}
.job-title .company, .degree-header .institution {{ font-weight: 600; }}
.job-title .position, .degree-header .area {{ font-weight: 400; color: var(--fg-1); }}
.job-title .sep, .degree-header .sep {{ color: var(--fg-2); margin-right: 2pt; }}
.dates {{ font-size: 9.5pt; color: var(--fg-2); font-variant-numeric: tabular-nums; white-space: nowrap; flex-shrink: 0; }}
.job-summary {{ font-size: 10pt; color: var(--fg-1); font-style: italic; margin-bottom: 3pt; }}
.highlights {{ list-style: disc; padding-left: 16pt; margin-top: 2pt; }}
.highlights li {{ margin-bottom: 2pt; }}

.skills .skill-line {{ line-height: 1.6; }}
.skill-groups {{ list-style: none; padding-left: 0; }}
.skill-groups li {{ margin-bottom: 4pt; line-height: 1.5; }}
.skill-groups strong {{ font-weight: 600; }}
.certifications ul {{ list-style: disc; padding-left: 16pt; }}
.certifications li {{ margin-bottom: 2pt; }}

article {{ page-break-inside: avoid; }}
@media print {{ body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }} }}
"""


def _build_html(css_filename: str) -> str:
    """Copy the canonical classic.html skeleton, swapping only the CSS href.

    The Jinja2 resume-data contract stays byte-identical to the bundled personas,
    so `pdf_render.render_html_string` / `_inline_persona_css` consume it unchanged.
    """
    skeleton = _SKELETON_HTML.read_text(encoding="utf-8")
    return skeleton.replace('href="classic.css"', f'href="{css_filename}"')


def generate_companion(docx_path: str | Path, *, force: bool = False) -> tuple[Path, Path] | None:
    """Write `<stem>.html` + `<stem>.css` next to an uploaded persona `.docx`.

    Returns `(html_path, css_path)`, or None if the .docx can't be read (the
    caller falls back to the bundled Classic companion — no worse than before).
    Idempotent: skips regeneration when both companions exist and are newer than
    the source .docx, unless `force=True`. Best-effort by design — a failure here
    must never break upload or preview.
    """
    docx = Path(docx_path)
    html_path = docx.with_suffix(".html")
    css_path = docx.with_suffix(".css")
    sidecar = docx.with_name(docx.stem + ".persona.json")

    try:
        if (
            not force
            and html_path.exists()
            and css_path.exists()
            and html_path.stat().st_mtime >= docx.stat().st_mtime
        ):
            return html_path, css_path

        knobs = extract_persona_style(docx)
        css_path.write_text(_build_css(knobs), encoding="utf-8")
        html_path.write_text(_build_html(css_path.name), encoding="utf-8")
        sidecar.write_text(
            json.dumps({"layout_fidelity": knobs["layout_fidelity"]}, indent=2),
            encoding="utf-8",
        )
        logger.info(
            "Generated persona companion for %s (layout_fidelity=%s)",
            docx.name,
            knobs["layout_fidelity"],
        )
        return html_path, css_path
    except Exception as exc:  # best-effort; a companion failure must never break upload/preview
        logger.warning("Persona companion generation failed for %s: %s", docx, exc)
        return None
